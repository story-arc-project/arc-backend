# -*- coding: utf-8 -*-
"""
Career Keyword Analyzer v4.1 - Precision K-NN + Deep Crawler Edition
=====================================================================
키워드 기반 경력 분석 시스템 (Google Colab 전용)

[v4.1 변경 사항 — 선별 정밀화 + 방향성 강화]
  - 키워드당 선별 경험 수 축소: 5개 → 3개 (_KNN_K)
      → 스토리라인이 핵심 경험에 집중, 약한 근거의 꼬리 경험 원천 차단
  - 억지 매칭 방지 3중 장치:
      1) 상대 컷 강화 (top × 0.80 → 0.85) — 최상위 대비 애매한 후보 탈락
      2) 프롬프트 규칙: 2단계 이상 추론이 필요한 연결은 제외,
         개수를 채우기 위한 포함 금지, low는 직접 인용 근거 1개 이상일 때만
      3) K-NN 선별 경험이라도 근거 불충분 시 LLM이 제외 가능하도록 명시
  - F_improvement_guide 방향성 강화: overall_direction 신설
      (현재 프로필 진단 → 단기 3개월 / 중기 1년 실행 방향 → 최우선 보강 키워드)
      각 개선 항목에 priority(높음/중간/낮음) 부여
  - print_result_summary: 종합 방향성 섹션 출력 추가

[v4.0 변경 사항 — K-NN 매칭 엔진 완성]
  - 하이브리드 스코어링: 코사인 유사도(의미) + 문자 바이그램 어휘 유사도(표면)
      score = 0.80 × cosine + 0.20 × lexical  (한국어 비정형 텍스트 최적화)
  - 멀티 팩싯 쿼리 확장: 키워드 → {정의, 행동 표현, 동의어, 성과 표현} 4개 관점으로
      각각 임베딩 후 센트로이드(평균 벡터) 사용 → 단일 문장 확장 대비 재현율 향상
  - 비대칭 태스크 타입: 경험 임베딩은 RETRIEVAL_DOCUMENT, 키워드는 RETRIEVAL_QUERY
  - 적응형 컷오프: 절대 임계값 + 최고점 대비 상대 컷(top × 0.80) 이중 필터
      → 전반적으로 유사도가 높은 풀에서 노이즈 유입 차단
  - MMR 다양성 재랭킹: 유사한 경험끼리 중복 선별되는 문제 해결 (λ=0.75)
  - 키워드 확장 결과 캐시: analyze_keywords 재호출 시 LLM/임베딩 호출 생략
  - 임베딩 벡터 단위 정규화 저장 → 코사인 계산 = 내적 (고속화)
  - 작동 임베딩 설정 캐시 실사용 (v3.0은 캐시만 하고 사용하지 않던 버그 수정)
  - explain_knn(): 키워드-경험 전체 매칭 점수 진단 유틸리티

[v4.0 변경 사항 — 딥 크롤러 완성]
  - 병렬 fetch: ThreadPoolExecutor 기반 동시 다운로드 (기본 6 워커) → 크롤링 수 배 고속화
  - 우선순위 큐: URL 경로의 커리어 관련 토큰(portfolio/project/이력/활동 등)으로
      방문 순서 결정 → 페이지 한도 내에서 정보 밀도 최대화
  - sitemap.xml 자동 발견·시딩 (sitemap index 1단계 추적 포함)
  - URL 정규화: 추적 파라미터(utm_*/fbclid/gclid 등) 제거 + 쿼리 정렬 + fragment 제거
      → 동일 페이지 중복 방문 원천 차단
  - 문자셋 자동 감지: HTTP 헤더 → meta charset → utf-8 → cp949/euc-kr 순차 폴백
  - fetch 네트워크 재시도 (지수 백오프)
  - 콘텐츠 지문(해시) 기반 중복 페이지 제거
  - 보일러플레이트 제거: 60% 이상 페이지에서 반복되는 짧은 줄(내비/푸터) 자동 삭제
  - JSON-LD(application/ld+json) 구조화 데이터 추출 (Person/Article 등)
  - Notion 서브페이지 자동 발견·크롤 (최대 5개, 깊이 1)
  - Playwright Notion 렌더링: 높이 안정화까지 반복 스크롤 (lazy-load 완전 로딩)
  - 페이지 한도 도달 후 불필요한 fetch를 하고 버리던 낭비 제거
  - _resolve_input() 단일 URL 분기에서 크롤 결과를 버리던 버그 수정
  - JS 템플릿 리터럴 필터 수정("${{" → "${") 및 og 메타 정규식 정리
  - .docx 파일 지원 (python-docx 설치 시)

[v3.0 기능 유지]
  - 하이브리드 파이프라인 (경험 10개 미만 → LLM 직접 / 이상 → K-NN 사전 필터링)
  - parse_careers() 시점 임베딩 사전 생성·캐시
  - coverage_percent 실측 계산
  - raw_content 절삭: LLM 모드 2,000자 / K-NN 모드 전체

[v2.1 기능 유지]
  - 복합 입력 지원 (URL + 파일경로 + 텍스트 혼합)
  - 딥 크롤링 (최대 30페이지 + 5 PDF)
  - 중복 경력 자동 제거 (MD5 ID 기반)
  - merge_careers() 병합 유틸리티

사용법:
  # 1. 경력 파싱
  careers = parse_careers("https://example.com/portfolio")
  careers = parse_careers("직접 텍스트 입력...")
  careers = parse_careers('''
      동아리소개: https://wafflestudio.com/
      인턴: 엔젠바이오 2023.7월~8월, 유전자 검사 시각화 프로그램 개발
  ''')

  # 2. 소스 병합
  more    = parse_careers("추가 텍스트...")
  careers = merge_careers(careers, more)

  # 3. 키워드 분석 (경험 수에 따라 자동 라우팅)
  result = analyze_keywords(["리더십", "문제해결"], careers)
  result = analyze_keywords(["협업"], careers, target="스타트업 PM")

  # 4. 추천 키워드 확인 / K-NN 진단
  show_keywords()
  explain_knn("리더십", careers)   # 전체 경험의 매칭 점수 확인
"""

import json
import re
import os
import io
import time
import heapq
import hashlib
import threading
import urllib.request
import urllib.parse
from html.parser import HTMLParser
from datetime import datetime
from collections import deque, Counter
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import Optional

from google import genai
from google.genai import types

try:
    import pypdf
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False
    print("WARNING: pip install pypdf  → PDF 파싱 기능 비활성화")

try:
    import docx as _docx  # python-docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from playwright.sync_api import sync_playwright, TimeoutError as PWTimeoutError
    HAS_PLAYWRIGHT = True
except ImportError:
    HAS_PLAYWRIGHT = False
    print("INFO: pip install playwright && playwright install chromium  → SPA/Notion 완전 렌더링 비활성화")


# ══════════════════════════════════════════════════════════════════
# 1. Config
# ══════════════════════════════════════════════════════════════════
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")  # ← 여기에 API 키 입력
client = None

_ANALYSIS_MODEL  = "gemini-2.5-flash"
_EMBEDDING_MODEL = "gemini-embedding-001"

# ── 크롤러 설정 ──────────────────────────────────────────────────
_MAX_PAGES        = 30    # 크롤링 최대 페이지 수
_MAX_PDFS         = 5     # 크롤링 최대 PDF 수
_FETCH_TIMEOUT    = 15    # 개별 요청 타임아웃 (초)
_MAX_CONTENT_MB   = 1     # 개별 응답 최대 크기 (MB)
_FETCH_RETRIES    = 2     # 네트워크 실패 시 fetch 재시도 횟수
_CRAWL_WORKERS    = 6     # 병렬 fetch 스레드 수
_BOILERPLATE_RATIO = 0.6  # 이 비율 이상 페이지에서 반복되는 줄은 내비/푸터로 간주해 제거
_NOTION_MAX_SUBPAGES = 5  # Notion 서브페이지 자동 크롤 한도

# ── LLM 재시도 설정 ─────────────────────────────────────────────
_MAX_RETRIES    = 4
_RETRY_BASE_SEC = 5

# ── K-NN 하이브리드 설정 ─────────────────────────────────────────
_KNN_THRESHOLD_COUNT = 10    # 이 수 이상이면 K-NN 파이프라인 활성화
_KNN_K               = 3     # 키워드당 선별할 최대 경험 수 (v4.1: 5→3, 핵심 경험 집중)
_KNN_SIM_THRESHOLD   = 0.40  # 하이브리드 점수 절대 최소 기준값
#   ↑ 조정 가이드:
#     0.35 미만 → 노이즈 경험 섞임
#     0.45 초과 → 관련 경험 오탈락 위험
#     한국어 비정형 텍스트 기준 0.40이 보수적 기본값
_KNN_COSINE_WEIGHT   = 0.80  # 하이브리드 점수 내 코사인(의미) 가중치
_KNN_LEXICAL_WEIGHT  = 0.20  # 하이브리드 점수 내 어휘(표면) 가중치
_KNN_RELATIVE_CUT    = 0.85  # 최고점 대비 상대 컷 (v4.1: 0.80→0.85, 억지 매칭 방지 강화)
_KNN_MMR_LAMBDA      = 0.75  # MMR 재랭킹: 1.0=관련도만, 0.0=다양성만

_EMB_TASK_DOC   = "RETRIEVAL_DOCUMENT"  # 경험(문서) 임베딩 태스크
_EMB_TASK_QUERY = "RETRIEVAL_QUERY"    # 키워드(질의) 임베딩 태스크

# 크롤링 결과 최소 유효 글자 수 (이하면 경고 및 힌트 보강)
_MIN_CRAWL_CHARS = 500

# 크롤링에서 항상 제외할 URL 패턴 (404 노이즈 차단)
_SKIP_URL_PATTERNS = (
    "cdn-cgi/",          # Cloudflare 내부 경로
    "/wp-json/",         # WordPress REST API
    "/feed/",            # RSS 피드
    "/wp-admin/",
    "/wp-login.php",
    "/xmlrpc.php",
    "/tag/", "/tags/",   # 태그 아카이브 (중복 콘텐츠)
    "/category/",        # 카테고리 아카이브
    "/page/",            # 페이지네이션 아카이브
    "/search",           # 검색 결과
    "/login", "/signin", "/signup", "/cart", "/checkout",
)
# 리소스 파일 확장자 — 별도 정규식으로 처리 (_should_skip_url 참조)
_SKIP_RESOURCE_EXTS = re.compile(
    r"\.(css|js|ts|map|woff2?|ttf|eot|otf|ico|png|jpe?g|gif|svg|webp"
    r"|mp4|webm|mp3|wav|ogg|zip|gz|tar|rar|7z|exe|dmg|apk"
    r"|xml|rss|atom)([?#]|$)",
    re.IGNORECASE,
)
# ※ .pdf는 리소스가 아니라 콘텐츠이므로 위 목록에서 제외 (별도 파싱 경로)

# URL 추적 파라미터 — 동일 페이지 중복 방문의 주범이므로 정규화 시 제거
_TRACKING_PARAM = re.compile(
    r"^(utm_\w+|fbclid|gclid|dclid|msclkid|igshid|mc_cid|mc_eid"
    r"|ref|ref_src|source|_ga|_gl|spm|share_token)$",
    re.IGNORECASE,
)

# URL 경로 우선순위 토큰 — 커리어 정보 밀도가 높은 경로를 먼저 방문
_URL_PRIORITY_TOKENS = {
    # 영문
    "portfolio": 6, "project": 6, "resume": 6, "career": 6,
    "cv": 5, "experience": 5, "intern": 5, "award": 5, "achievement": 5,
    "about": 4, "profile": 4, "activity": 4, "research": 4,
    "publication": 4, "certificate": 4,
    "work": 3, "education": 3, "team": 3, "member": 3,
    "blog": 2, "post": 2, "article": 2,
    # 국문 (URL에 퍼센트 인코딩된 경우 unquote 후 매칭)
    "포트폴리오": 6, "프로젝트": 6, "이력": 6, "경력": 6,
    "수상": 5, "활동": 5, "인턴": 5,
    "소개": 4, "연구": 4, "자격": 4,
    "학력": 3, "팀": 3,
}

# 크롤 결과 캐시 — 동일 URL 이중 크롤 방지 (프로세스 생존 동안 유지)
_URL_CRAWL_CACHE: dict[str, str] = {}

KEYWORD_CATEGORIES = {
    # ── 역량: "무엇을 할 수 있는가" — 행동 중심 (18개) ─────────────
    # 데이터분석능력 → 분석력에 통합 / 갈등조정 → 팀워크에 포함
    "역량": [
        "리더십",      "팀워크",       "커뮤니케이션", "문제해결력",
        "기획력",      "실행력",       "분석력",       "창의력",
        "적응력",      "학습능력",     "설득력",       "조직관리",
        "프로젝트관리","전략적사고",   "위기관리",     "멘토링코칭",
        "자기관리",    "협상력",
    ],
    # ── 직무: 조직 안에서 맡는 기능적 역할 — 세분화 버전 (29개) ────
    "직무": [
        # 기획/전략
        "사업기획",        "전략기획",        "서비스기획",       "신사업개발",
        # 마케팅
        "브랜드마케팅",    "디지털마케팅",    "퍼포먼스마케팅",
        "콘텐츠마케팅",    "CRM마케팅",
        # 영업/BD
        "B2B영업",         "B2C영업",         "사업개발BD",
        # 데이터/개발
        "데이터분석",      "AI_ML엔지니어링", "프론트엔드개발",
        "백엔드개발",      "DevOps인프라",
        # 디자인
        "UX리서치",        "UI디자인",        "브랜드그래픽디자인",
        # HR
        "채용리크루팅",    "조직개발HRD",
        # 재무
        "재무기획FPA",     "IR투자분석",
        # PM/운영/고객
        "프로덕트PM",      "고객경험CX",
        # 기타 전문직무
        "PR홍보",          "법률",            "스타트업(창업)",
    ],
    # ── 직종: 전문직 커리어 트랙 (20개) ────────────────────────────
    "직종": [
        "IT개발직",   "PM기획직",     "법률직",       "의료직",
        "금융직",     "컨설팅직",     "학술연구직",   "교육직",
        "디자인직",   "회계세무직",   "언론미디어직", "건축도시직",
        "공학기술직", "공공행정직",   "사회복지직",   "창업가",
        "데이터AI직", "영업세일즈직", "예술창작직",   "마케팅광고직",
    ],
    # ── 진로 관련 가치관: "무엇을 중요하게 여기는가" (20개) ─────────
    "진로관련가치관": [
        "성취지향적",    "사회적기여",    "자율성추구",    "안정성추구",
        "도전지향적",    "영향력추구",    "전문성추구",    "창의적표현",
        "협력중심적",    "윤리가치기반",  "균형추구",      "성장지향적",
        "다양성존중",    "글로벌지향",    "혁신추구",      "관계중심적",
        "데이터기반신뢰","투명성중시",    "지속가능성",    "주도적",
    ],
}


# ══════════════════════════════════════════════════════════════════
# 2. LLM 호출 헬퍼
# ══════════════════════════════════════════════════════════════════
def _init_client():
    global client
    if client is None:
        if not GEMINI_API_KEY:
            raise ValueError("❌ GEMINI_API_KEY가 설정되지 않았습니다.")
        client = genai.Client(api_key=GEMINI_API_KEY)
    return client


def _is_rate_limit(e: Exception) -> bool:
    s = str(e).lower()
    return any(k in s for k in ("429", "quota", "rate_limit", "rate limit", "resource_exhausted"))


def _call_with_retry(fn, *args, **kwargs):
    for attempt in range(_MAX_RETRIES):
        try:
            return fn(*args, **kwargs)
        except Exception as e:
            if _is_rate_limit(e) and attempt < _MAX_RETRIES - 1:
                wait = _RETRY_BASE_SEC * (2 ** attempt)
                print(f"  [Rate Limit] {wait}초 후 재시도 ({attempt + 1}/{_MAX_RETRIES - 1})...", flush=True)
                time.sleep(wait)
            else:
                raise


def _call_model(system_prompt: str, user_prompt: str, max_tokens: int = 65536) -> dict:
    _init_client()
    raw_text = ""
    try:
        def _do():
            return client.models.generate_content(
                model=_ANALYSIS_MODEL,
                contents=user_prompt,
                config=types.GenerateContentConfig(
                    system_instruction=system_prompt if system_prompt else None,
                    temperature=0.1,
                    max_output_tokens=max_tokens,
                    response_mime_type="application/json",
                ),
            )
        resp = _call_with_retry(_do)
        raw_text = resp.text.strip()
        return json.loads(_clean_json_response(raw_text))
    except json.JSONDecodeError as e:
        recovered = _try_recover_truncated_json(raw_text)
        if recovered:
            print(f"  [JSON 복구 성공] 절단된 응답 복구 완료", flush=True)
            return recovered
        print(f"  [JSON 복구 실패] 원인: {e.msg} / 위치: char {e.pos}", flush=True)
        return {"status": "error", "message": f"JSON parse failed: {e.msg}", "raw_response": raw_text[:2000]}
    except Exception as e:
        return {"status": "error", "message": str(e)}


def _try_recover_truncated_json(raw_text: str) -> dict | None:
    """토큰 한도로 절단된 JSON 복구 시도. 실패 시 None 반환."""
    text = _clean_json_response(raw_text)
    stack, in_str, escaped, last_valid_pos = [], False, False, 0

    for i, ch in enumerate(text):
        if escaped:
            escaped = False
            continue
        if ch == "\\" and in_str:
            escaped = True
            continue
        if ch == '"':
            in_str = not in_str
        elif not in_str:
            if ch in "{[":
                stack.append(ch)
            elif ch in "}]":
                if stack:
                    stack.pop()
                if not stack:
                    last_valid_pos = i + 1

    if last_valid_pos > 0:
        try:
            return json.loads(text[:last_valid_pos])
        except Exception:
            pass

    if stack:
        closing = {"[": "]", "{": "}"}
        tail = "".join(closing.get(c, "") for c in reversed(stack))
        candidate = text.rstrip().rstrip(",") + ('"' if in_str else "") + tail
        try:
            return json.loads(candidate)
        except Exception:
            pass

    return None


def _sanitize_text(text: str) -> str:
    """
    LLM에 전달하기 전 텍스트를 정제합니다.
    JSON 문자열을 깨뜨리는 제어문자와 이스케이프 문제를 방지합니다.
    - 널 바이트 / 제어문자 제거 (탭·줄바꿈 제외)
    - 단독 역슬래시를 공백으로 치환
    - 연속 공백·줄바꿈 정리
    """
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    text = re.sub(r'\\(?![\"\'/nrtbfu])', " ", text)
    text = re.sub(r" {3,}", "  ", text)
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    return text.strip()


def _clean_json_response(raw_text: str) -> str:
    raw_text = raw_text.strip()
    m = re.search(r"```(?:json)?\s*\n?(.*?)\n?\s*```", raw_text, re.DOTALL | re.IGNORECASE)
    if m:
        return m.group(1).strip()
    depth, start = 0, -1
    for i, ch in enumerate(raw_text):
        if ch == "{":
            if depth == 0:
                start = i
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0 and start != -1:
                return raw_text[start:i + 1].strip()
    return raw_text


# ══════════════════════════════════════════════════════════════════
# 3. 임베딩 엔진 & 유사도 (v4.0 강화)
# ══════════════════════════════════════════════════════════════════
def _extract_vector(result) -> list | None:
    """embed_content 응답에서 벡터 추출. SDK 버전마다 구조가 다름."""
    if hasattr(result, "embeddings") and result.embeddings:
        emb = result.embeddings[0]
        if hasattr(emb, "values") and emb.values:
            return list(emb.values)
    if hasattr(result, "embedding") and result.embedding:
        emb = result.embedding
        if hasattr(emb, "values") and emb.values:
            return list(emb.values)
    return None


def _unit_normalize(vec: list) -> list:
    """벡터를 단위 길이로 정규화. 정규화 벡터끼리 코사인 = 내적 → 고속."""
    norm = sum(x * x for x in vec) ** 0.5
    if norm == 0:
        return vec
    return [x / norm for x in vec]


# 첫 성공 조합 캐시 — {"model": str, "param": "contents"|"content", "with_config": bool}
_WORKING_EMB_CONFIG: dict | None = None
_EMB_CONFIG_LOCK = threading.Lock()


def _build_embed_kwargs(param: str, text: str, with_config: bool, task_type: str) -> dict:
    kwargs = {param: text}
    if with_config:
        try:
            kwargs["config"] = types.EmbedContentConfig(task_type=task_type)
        except Exception:
            pass
    return kwargs


def _set_working_embedding_config(model: str, param: str, with_config: bool) -> None:
    global _WORKING_EMB_CONFIG
    with _EMB_CONFIG_LOCK:
        if _WORKING_EMB_CONFIG is None:
            _WORKING_EMB_CONFIG = {"model": model, "param": param, "with_config": with_config}
            print(f"  [임베딩] 작동 설정 캐시됨 → model={model}, param={param}, config={with_config}",
                  flush=True)


def get_embedding(text: str, task_type: str = _EMB_TASK_DOC) -> Optional[list]:
    """
    텍스트를 단위 정규화된 임베딩 벡터로 변환.

    v4.0 개선:
      - 작동 조합 캐시를 실제로 사용 (첫 성공 이후 단일 호출로 처리)
      - task_type 지정 지원 (문서=RETRIEVAL_DOCUMENT / 질의=RETRIEVAL_QUERY)
        → 비대칭 임베딩으로 질의-문서 매칭 정확도 향상
      - Rate limit 시 지수 백오프 재시도
      - 반환 벡터는 단위 정규화 완료 상태

    실패 시 None 반환.
    """
    if not text or not text.strip():
        return None
    _init_client()
    truncated = text[:10000]

    # 1) 캐시된 작동 조합이 있으면 그것부터 사용
    cfg = _WORKING_EMB_CONFIG
    if cfg:
        try:
            kwargs = _build_embed_kwargs(cfg["param"], truncated, cfg["with_config"], task_type)
            result = _call_with_retry(client.models.embed_content, model=cfg["model"], **kwargs)
            vec = _extract_vector(result)
            if vec:
                return _unit_normalize(vec)
        except Exception as e:
            if _is_rate_limit(e):
                raise
            print(f"  [임베딩] 캐시 설정 실패 → 전체 조합 재탐색: {e}", flush=True)

    # 2) 전체 조합 탐색 (SDK 버전·모델명·파라미터명 차이 커버)
    attempts = [
        # (설명, 모델명, 파라미터명, config 사용 여부)
        ("신버전·gemini-embedding-001·contents+config", _EMBEDDING_MODEL,          "contents", True),
        ("신버전·gemini-embedding-001·contents",        _EMBEDDING_MODEL,          "contents", False),
        ("신버전·text-embedding-004·contents",          "text-embedding-004",      "contents", False),
        ("구버전·gemini-embedding-001·content",         _EMBEDDING_MODEL,          "content",  False),
        ("구버전·models/prefix·content",                f"models/{_EMBEDDING_MODEL}", "content", False),
        ("구버전·text-embedding-004·content",           "text-embedding-004",      "content",  False),
    ]

    for label, model, param, with_config in attempts:
        try:
            kwargs = _build_embed_kwargs(param, truncated, with_config, task_type)
            result = client.models.embed_content(model=model, **kwargs)
            vec = _extract_vector(result)
            if vec:
                _set_working_embedding_config(model, param, with_config)
                return _unit_normalize(vec)
            print(f"  [임베딩] {label} — 응답은 왔으나 벡터 없음 (구조: {type(result)})",
                  flush=True)
        except Exception as e:
            print(f"  [임베딩 실패] {label} — {type(e).__name__}: {e}", flush=True)

    return None


def test_embedding(sample_text: str = "리더십과 문제해결 역량을 보유한 개발자입니다.") -> None:
    """
    임베딩 API 진단 함수.
    실패 원인을 상세히 출력하고 작동하는 설정을 알려줍니다.
    """
    print("=" * 60)
    print("  임베딩 API 진단")
    print("=" * 60)
    _init_client()

    attempts = [
        ("신버전·gemini-embedding-001·contents+config", _EMBEDDING_MODEL,          "contents", True),
        ("신버전·gemini-embedding-001·contents",        _EMBEDDING_MODEL,          "contents", False),
        ("신버전·text-embedding-004·contents",          "text-embedding-004",      "contents", False),
        ("구버전·gemini-embedding-001·content",         _EMBEDDING_MODEL,          "content",  False),
        ("구버전·models/prefix·content",                f"models/{_EMBEDDING_MODEL}", "content", False),
        ("구버전·text-embedding-004·content",           "text-embedding-004",      "content",  False),
    ]

    success = False
    for label, model, param, with_config in attempts:
        print(f"\n  시도: {label}")
        try:
            kwargs = _build_embed_kwargs(param, sample_text, with_config, _EMB_TASK_DOC)
            result = client.models.embed_content(model=model, **kwargs)
            vec = _extract_vector(result)
            if vec:
                print(f"  ✅ 성공! 벡터 차원={len(vec)}, 첫 값={vec[0]:.6f}")
                print(f"     → 작동 설정: model={model!r}, param={param!r}, config={with_config}")
                success = True
                break
            else:
                print(f"  ⚠️  응답 왔으나 벡터 없음. result type={type(result)}")
                print(f"     result attrs={[a for a in dir(result) if not a.startswith('_')]}")
        except Exception as e:
            print(f"  ❌ {type(e).__name__}: {e}")

    print()
    if not success:
        print("  모든 조합 실패. 확인 사항:")
        print("  1. GEMINI_API_KEY가 올바른지 확인")
        print("  2. Google AI Studio에서 Embedding API 활성화 여부 확인")
        print("  3. google-genai 패키지 버전 확인: pip show google-genai")
        print("  4. 네트워크/방화벽으로 API 호출이 차단되지 않는지 확인")
    print("=" * 60)


def cosine_similarity(a: list, b: list) -> float:
    """두 벡터의 코사인 유사도 반환. 단위 정규화 벡터라면 내적과 동일."""
    if not a or not b or len(a) != len(b):
        return 0.0
    dot    = sum(x * y for x, y in zip(a, b))
    norm_a = sum(x * x for x in a) ** 0.5
    norm_b = sum(x * x for x in b) ** 0.5
    if norm_a == 0 or norm_b == 0:
        return 0.0
    return dot / (norm_a * norm_b)


def _char_bigrams(text: str) -> set:
    """공백 정규화 후 문자 2-gram 집합. 형태소 분석기 없이 한국어 표면 매칭."""
    text = re.sub(r"\s+", " ", text.lower()).strip()
    if len(text) < 2:
        return {text} if text else set()
    return {text[i:i + 2] for i in range(len(text) - 1)}


def lexical_containment(query_text: str, doc_text: str) -> float:
    """
    질의 바이그램이 문서에 포함된 비율 (0.0 ~ 1.0).
    Jaccard 대신 포함률(containment)을 사용 — 문서가 길어도 점수가 희석되지 않음.
    임베딩이 놓치는 표면 어휘 일치(고유명사·전문용어)를 보완합니다.
    """
    q = _char_bigrams(query_text)
    if not q:
        return 0.0
    d = _char_bigrams(doc_text[:8000])
    return len(q & d) / len(q)


# ══════════════════════════════════════════════════════════════════
# 4. 크롤러 (fetch 엔진 · HTML 파서 · Notion · 딥크롤러 · 파일 리더)
# ══════════════════════════════════════════════════════════════════
class _TextExtractor(HTMLParser):
    SKIP_TAGS = {"script", "style", "noscript", "meta", "link", "head", "svg", "iframe"}

    def __init__(self, base_url: str = ""):
        super().__init__()
        self._skip = 0
        self._parts: list[str] = []
        self._links: list[str] = []
        self._base = base_url

    def handle_starttag(self, tag, attrs):
        tl = tag.lower()
        if tl in self.SKIP_TAGS:
            self._skip += 1
        if tl == "a":
            href = dict(attrs).get("href", "")
            if href and not href.startswith(("javascript:", "mailto:", "tel:")):
                self._links.append(urllib.parse.urljoin(self._base, href))

    def handle_endtag(self, tag):
        if tag.lower() in self.SKIP_TAGS and self._skip > 0:
            self._skip -= 1

    def handle_data(self, data):
        if self._skip == 0:
            s = data.strip()
            if s:
                self._parts.append(s)

    @property
    def text(self) -> str:
        return "\n".join(self._parts)

    @property
    def links(self) -> list[str]:
        return self._links


def _is_spa(html: str) -> bool:
    """JavaScript 렌더링 기반 SPA 페이지 여부 감지."""
    visible = len(re.sub(r"<[^>]+>", " ", html).strip())
    return visible < 2000 and html.lower().count("<script") >= 2


def _extract_links_from_raw_html(raw_html: str, base_url: str) -> list[str]:
    """href/src/data-href/data-src 속성에서 링크를 직접 추출 (SPA 대응).
    JS 템플릿 리터럴(${...})과 리소스 파일은 자동 제외."""
    found = []
    for m in re.finditer(
        r'(?:href|data-href|data-src)\s*=\s*["\']([^"\'\s]+)["\']',
        raw_html, re.IGNORECASE
    ):
        href = m.group(1).strip()
        # JS 템플릿 리터럴 제외: ${link.href}, ${project.page} 등
        if "${" in href:
            continue
        if href.startswith(("javascript:", "mailto:", "tel:", "#", "data:")):
            continue
        # 리소스 파일 확장자 제외
        if _SKIP_RESOURCE_EXTS.search(urllib.parse.urlparse(href).path):
            continue
        found.append(urllib.parse.urljoin(base_url, href))
    return list(dict.fromkeys(found))


def _extract_json_ld(html: str) -> str:
    """
    application/ld+json 구조화 데이터에서 커리어 관련 필드 추출.
    Person / Organization / Article / CreativeWork 스키마의 텍스트 필드를 수집.
    """
    _FIELDS = ("name", "jobTitle", "worksFor", "alumniOf", "description",
               "headline", "about", "award", "knowsAbout", "affiliation",
               "author", "datePublished")
    parts: list[str] = []

    def _walk(node):
        if isinstance(node, dict):
            for key in _FIELDS:
                v = node.get(key)
                if isinstance(v, str) and v.strip():
                    parts.append(f"{key}: {v.strip()}")
                elif isinstance(v, dict) and isinstance(v.get("name"), str):
                    parts.append(f"{key}: {v['name']}")
                elif isinstance(v, list):
                    names = [x if isinstance(x, str) else x.get("name", "")
                             for x in v if isinstance(x, (str, dict))]
                    names = [n for n in names if n]
                    if names:
                        parts.append(f"{key}: {', '.join(names[:10])}")
            for v in node.values():
                if isinstance(v, (dict, list)):
                    _walk(v)
        elif isinstance(node, list):
            for item in node:
                _walk(item)

    for m in re.finditer(
        r'<script[^>]+type=["\']application/ld\+json["\'][^>]*>(.*?)</script>',
        html, re.DOTALL | re.IGNORECASE
    ):
        try:
            _walk(json.loads(m.group(1).strip()))
        except (json.JSONDecodeError, RecursionError):
            continue

    # 중복 제거, 상한
    seen, out = set(), []
    for p in parts:
        if p not in seen:
            seen.add(p)
            out.append(p)
    return "\n".join(out[:40])


# ── URL 정규화 · 우선순위 · 스킵 판단 ────────────────────────────

def _normalize_url(url: str) -> str:
    """
    URL 정규화: 소문자 호스트, fragment 제거, 추적 파라미터 제거,
    쿼리 정렬, 후행 슬래시 제거 → 동일 페이지 중복 방문 차단.
    """
    try:
        p = urllib.parse.urlsplit(url)
    except ValueError:
        return url
    query_pairs = [
        (k, v) for k, v in urllib.parse.parse_qsl(p.query, keep_blank_values=False)
        if not _TRACKING_PARAM.match(k)
    ]
    query = urllib.parse.urlencode(sorted(query_pairs))
    path = p.path.rstrip("/")
    return urllib.parse.urlunsplit((p.scheme.lower(), p.netloc.lower(), path, query, ""))


def _url_priority(url: str) -> int:
    """URL 경로의 커리어 관련 토큰 가중치 합산. 높을수록 먼저 방문."""
    path = urllib.parse.unquote(urllib.parse.urlparse(url).path.lower())
    return sum(w for tok, w in _URL_PRIORITY_TOKENS.items() if tok in path)


def _should_skip_url(url: str) -> bool:
    """노이즈성 URL 여부 판단. True면 크롤링 대상에서 제외."""
    lower = url.lower()
    if any(pat in lower for pat in _SKIP_URL_PATTERNS):
        return True
    path = urllib.parse.urlparse(url).path
    if _SKIP_RESOURCE_EXTS.search(path):
        return True
    return False


def _same_origin(a: str, b: str) -> bool:
    return urllib.parse.urlparse(a).netloc == urllib.parse.urlparse(b).netloc


def _is_pdf_url(url: str, data: bytes | None = None) -> bool:
    if url.lower().split("?")[0].endswith(".pdf"):
        return True
    return bool(data and data[:4] == b"%PDF")


# ── fetch 엔진 (재시도 + 문자셋 감지) ────────────────────────────

def _fetch(url: str, quiet: bool = False) -> tuple[bytes | None, str]:
    """
    URL을 fetch하여 (바이트, Content-Type 헤더) 반환.
    네트워크 오류 시 지수 백오프로 최대 _FETCH_RETRIES회 재시도.
    """
    last_err = None
    for attempt in range(_FETCH_RETRIES + 1):
        try:
            req = urllib.request.Request(
                url,
                headers={
                    "User-Agent": (
                        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                        "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0 Safari/537.36"
                    ),
                    "Accept": "text/html,application/xhtml+xml,application/pdf,*/*;q=0.8",
                    "Accept-Language": "ko-KR,ko;q=0.9,en-US;q=0.8",
                },
            )
            with urllib.request.urlopen(req, timeout=_FETCH_TIMEOUT) as resp:
                ctype = resp.headers.get("Content-Type", "")
                return resp.read(int(_MAX_CONTENT_MB * 1024 * 1024)), ctype
        except Exception as e:
            last_err = e
            # 404/403 등 명확한 HTTP 오류는 재시도 무의미
            if isinstance(e, urllib.error.HTTPError) and e.code in (400, 401, 403, 404, 410):
                break
            if attempt < _FETCH_RETRIES:
                time.sleep(1.5 * (2 ** attempt))
    if not quiet:
        print(f"    WARNING fetch failed [{url}]: {last_err}", flush=True)
    return None, ""


def _fetch_bytes(url: str) -> bytes | None:
    """하위 호환 래퍼: 바이트만 반환."""
    data, _ = _fetch(url)
    return data


def _decode_html(data: bytes, ctype: str = "") -> str:
    """
    문자셋 자동 감지 디코딩.
    HTTP 헤더 → meta charset → utf-8 → cp949 → euc-kr 순으로 시도.
    """
    candidates: list[str] = []
    m = re.search(r"charset=([\w-]+)", ctype, re.IGNORECASE)
    if m:
        candidates.append(m.group(1))
    head = data[:4096].decode("ascii", errors="ignore")
    m2 = re.search(r'charset\s*=\s*["\']?([\w-]+)', head, re.IGNORECASE)
    if m2:
        candidates.append(m2.group(1))
    candidates += ["utf-8", "cp949", "euc-kr"]
    for enc in candidates:
        try:
            return data.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
    return data.decode("utf-8", errors="replace")


def _parse_pdf_bytes(data: bytes) -> str:
    if not HAS_PYPDF:
        print("    WARNING: pypdf 미설치 — PDF 파싱 불가", flush=True)
        return ""
    try:
        reader = pypdf.PdfReader(io.BytesIO(data))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        print(f"    OK PDF parsed ({len(reader.pages)} pages, {len(text)} chars)", flush=True)
        return text
    except Exception as e:
        print(f"    WARNING PDF parse error: {e}", flush=True)
        return ""


def _parse_html_bytes(data: bytes, url: str, ctype: str = "") -> tuple[str, list[str], str]:
    """HTML 바이트를 파싱하여 (텍스트, 링크목록, 원본HTML) 반환."""
    html = _decode_html(data, ctype)
    parser = _TextExtractor(base_url=url)
    try:
        parser.feed(html)
    except Exception:
        pass
    return parser.text, parser.links, html


def _content_fingerprint(text: str) -> str:
    """공백/숫자 정규화 후 콘텐츠 지문 생성 — 실질 중복 페이지 감지."""
    normalized = re.sub(r"[\s\d]+", " ", text).strip().lower()
    return hashlib.md5(normalized.encode("utf-8")).hexdigest()


def _fetch_sitemap_urls(start_url: str, limit: int = 80) -> list[str]:
    """
    sitemap.xml 자동 발견. sitemap index인 경우 하위 sitemap을 1단계 추적.
    실패해도 조용히 빈 리스트 반환 (크롤링은 링크 추적만으로 진행).
    """
    parsed = urllib.parse.urlparse(start_url)
    root = f"{parsed.scheme}://{parsed.netloc}"
    urls: list[str] = []

    def _locs(data: bytes) -> list[str]:
        return [l.decode("utf-8", "ignore").strip()
                for l in re.findall(rb"<loc>\s*([^<\s]+)\s*</loc>", data)]

    for path in ("/sitemap.xml", "/sitemap_index.xml"):
        data, _ = _fetch(root + path, quiet=True)
        if not data or b"<" not in data[:200]:
            continue
        locs = _locs(data)
        # sitemap index → 하위 sitemap 1단계 추적 (최대 3개)
        child_maps = [l for l in locs if l.lower().split("?")[0].endswith(".xml")]
        page_locs  = [l for l in locs if l not in child_maps]
        for child in child_maps[:3]:
            cdata, _ = _fetch(child, quiet=True)
            if cdata:
                page_locs.extend(_locs(cdata))
        urls = [u for u in page_locs if u.startswith("http")][:limit]
        if urls:
            print(f"    sitemap 발견: {len(urls)}개 URL 시딩", flush=True)
            break
    return urls


def _strip_boilerplate(pages: list[str]) -> list[str]:
    """
    여러 페이지에서 반복 등장하는 짧은 줄(내비게이션·푸터·쿠키 배너)을 제거.
    _BOILERPLATE_RATIO 이상 페이지에 나타나는 120자 미만 줄이 대상.
    페이지가 3개 미만이면 판단 근거 부족으로 원본 유지.
    """
    if len(pages) < 3:
        return pages
    line_doc_freq: Counter = Counter()
    for text in pages:
        for ln in {l.strip() for l in text.splitlines() if l.strip()}:
            line_doc_freq[ln] += 1
    min_docs = max(3, int(len(pages) * _BOILERPLATE_RATIO))
    boiler = {ln for ln, c in line_doc_freq.items() if c >= min_docs and len(ln) < 120}
    if not boiler:
        return pages
    cleaned = []
    removed_total = 0
    for text in pages:
        kept_lines = [l for l in text.splitlines() if l.strip() not in boiler]
        removed_total += len(text.splitlines()) - len(kept_lines)
        cleaned.append("\n".join(kept_lines))
    print(f"    보일러플레이트 제거: 반복 줄 {len(boiler)}종 / 총 {removed_total}줄 삭제", flush=True)
    return cleaned


# ── Notion 전용 크롤러 ────────────────────────────────────────────

def _is_notion_url(url: str) -> bool:
    host = urllib.parse.urlparse(url).netloc.lower()
    return "notion.site" in host or "notion.so" in host


def _parse_notion_blocks(blocks: dict) -> str:
    lines: list[str] = []

    def _rich_text_to_str(rt) -> str:
        if not isinstance(rt, list):
            return ""
        parts = []
        for chunk in rt:
            if isinstance(chunk, list) and chunk:
                parts.append(str(chunk[0]))
            elif isinstance(chunk, str):
                parts.append(chunk)
        return "".join(parts).strip()

    TEXT_PROPS = ["title", "caption", "description"]
    SKIP_TYPES = {"image", "video", "file", "audio", "pdf", "embed", "bookmark",
                  "divider", "table_of_contents", "breadcrumb", "unsupported"}

    for block_id, block_wrapper in blocks.items():
        val = block_wrapper.get("value", {})
        if not isinstance(val, dict):
            continue
        btype = val.get("type", "")
        if btype in SKIP_TYPES:
            continue
        props = val.get("properties", {})
        if not isinstance(props, dict):
            continue
        for prop_key in TEXT_PROPS:
            if prop_key in props:
                text = _rich_text_to_str(props[prop_key])
                if text:
                    lines.append(text)
                break

    return "\n".join(lines)


def _extract_from_next_data(html: str) -> str:
    m = re.search(
        r'<script[^>]+id=["\']__NEXT_DATA__["\'][^>]*>\s*(.*?)\s*</script>',
        html, re.DOTALL | re.IGNORECASE
    )
    if not m:
        return ""
    try:
        data = json.loads(m.group(1))
    except json.JSONDecodeError:
        return ""
    record_map = (
        data.get("props", {}).get("pageProps", {}).get("recordMap", {})
    )
    if not record_map:
        record_map = data.get("props", {}).get("pageProps", {})
    blocks = record_map.get("block", {})
    if not blocks:
        return ""
    text = _parse_notion_blocks(blocks)
    print(f"    __NEXT_DATA__ 파싱 성공: {len(blocks)}개 블록, {len(text)} chars", flush=True)
    return text


def _extract_og_meta(html: str, url: str) -> str:
    parts = []
    patterns = [
        (r'<title[^>]*>(.*?)</title>', "페이지 제목"),
        (r'<meta[^>]+property=["\']og:title["\'][^>]+content=["\'](.*?)["\']', "OG 제목"),
        (r'<meta[^>]+property=["\']og:description["\'][^>]+content=["\'](.*?)["\']', "OG 설명"),
        (r'<meta[^>]+name=["\']description["\'][^>]+content=["\'](.*?)["\']', "메타 설명"),
    ]
    seen = set()
    for pattern, label in patterns:
        m = re.search(pattern, html, re.IGNORECASE | re.DOTALL)
        if m:
            val = m.group(1).strip()
            val = val.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")\
                     .replace("&#39;", "'").replace("&quot;", '"')
            if val and val not in seen:
                parts.append(f"{label}: {val}")
                seen.add(val)
    path = urllib.parse.urlparse(url).path.rstrip("/")
    slug = path.split("/")[-1]
    slug_clean = re.sub(r"[-_]", " ", re.sub(r"-[0-9a-f]{20,}$", "", slug)).strip()
    if slug_clean:
        parts.append(f"URL 슬러그(페이지명 힌트): {slug_clean}")
    return "\n".join(parts)


_NOTION_PAGE_ID = re.compile(r"[0-9a-f]{32}|[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}")


def _extract_notion_subpage_links(html: str, base_url: str) -> list[str]:
    """렌더링된 Notion HTML에서 동일 워크스페이스 서브페이지 링크 추출."""
    base_norm = _normalize_url(base_url)
    out = []
    for link in _extract_links_from_raw_html(html, base_url):
        if not _is_notion_url(link):
            continue
        norm = _normalize_url(link)
        if norm == base_norm:
            continue
        # Notion 페이지 ID가 경로에 있어야 실제 페이지로 판단
        if not _NOTION_PAGE_ID.search(urllib.parse.urlparse(link).path.replace("-", "")):
            if not _NOTION_PAGE_ID.search(urllib.parse.urlparse(link).path):
                continue
        if norm not in out:
            out.append(norm)
    return out


def crawl_notion_page(url: str, _depth: int = 0) -> str:
    """
    Notion 페이지 전용 크롤러.
    Playwright 설치 시: 높이 안정화까지 반복 스크롤 후 완전 렌더링 (권장).
    미설치 시: __NEXT_DATA__ + OG 메타 + HTML 텍스트 순으로 폴백.
    v4.0: 서브페이지 자동 발견·크롤 (최대 _NOTION_MAX_SUBPAGES개, 깊이 1).
    동일 URL은 캐시에서 즉시 반환.
    """
    cache_key = _normalize_url(url)
    if cache_key in _URL_CRAWL_CACHE:
        print(f"  [Notion 크롤러] 캐시 히트 → 재크롤 생략: {url}", flush=True)
        return _URL_CRAWL_CACHE[cache_key]
    print(f"  [Notion 크롤러] 시작 (depth={_depth}): {url}", flush=True)
    collected_parts: list[str] = []

    # Playwright 우선 시도 — Notion은 JS 렌더링 없이는 콘텐츠 대부분 누락
    html = ""
    with _PlaywrightSession() as pw:
        if pw.available:
            rendered = pw.render_notion(url)
            if rendered:
                html = rendered
                print(f"    Playwright HTML 사용 ({len(html)} chars)", flush=True)

    # Playwright 실패 또는 미설치 시 일반 fetch로 폴백
    if not html:
        raw_bytes, ctype = _fetch(url)
        if raw_bytes:
            html = _decode_html(raw_bytes, ctype)
        print(f"    일반 fetch HTML 사용 ({len(html)} chars)", flush=True)

    if html:
        next_data_text = _extract_from_next_data(html)
        if next_data_text.strip():
            collected_parts.append(f"[Notion 블록 콘텐츠: {url}]\n{next_data_text}")

    og_text = _extract_og_meta(html, url)
    if og_text.strip():
        collected_parts.append(f"[Notion 메타정보: {url}]\n{og_text}")

    if html:
        parser = _TextExtractor(base_url=url)
        try:
            parser.feed(html)
        except Exception:
            pass
        plain_text = parser.text.strip()
        if plain_text and len(plain_text) > 100:
            collected_parts.append(f"[Notion HTML 텍스트: {url}]\n{plain_text}")

    # ── v4.0: 서브페이지 자동 크롤 (깊이 1 한정) ─────────────────
    if html and _depth == 0:
        sub_links = _extract_notion_subpage_links(html, url)[:_NOTION_MAX_SUBPAGES]
        if sub_links:
            print(f"    Notion 서브페이지 {len(sub_links)}개 발견 → 크롤 시작", flush=True)
        for sub in sub_links:
            sub_content = crawl_notion_page(sub, _depth=1)
            if sub_content and len(sub_content) > 200:
                collected_parts.append(f"[Notion 서브페이지: {sub}]\n{sub_content}")

    combined = "\n\n".join(collected_parts).strip()
    print(f"    추출 결과: {len(combined)} chars", flush=True)

    if len(combined) < _MIN_CRAWL_CHARS and _depth == 0:
        print(f"    WARNING: 결과 빈약 → URL 힌트 보강", flush=True)
        path = urllib.parse.urlparse(url).path.rstrip("/")
        slug_raw = path.split("/")[-1]
        slug_clean = re.sub(r"[-_]", " ", re.sub(r"-[0-9a-f]{20,}$", "", slug_raw)).strip()
        hint = (
            f"[Notion 크롤링 제한 안내]\n"
            f"Notion 페이지 URL: {url}\n"
            f"페이지 제목/경로명: {slug_clean}\n"
            f"Notion은 JavaScript 렌더링 기반이라 완전한 크롤링이 제한될 수 있습니다.\n"
            f"위 페이지명과 수집된 정보를 바탕으로 최대한 추론하되, "
            f"추론 불가 항목은 반드시 빈 배열로 반환하십시오.\n"
        )
        combined = hint + "\n\n" + combined

    print(f"  [Notion 크롤러 완료] 총 {len(combined)} chars", flush=True)
    _URL_CRAWL_CACHE[cache_key] = combined
    return combined


# ══════════════════════════════════════════════════════════════════
# 4-1. Playwright 세션 (SPA / Notion 렌더링)
# ══════════════════════════════════════════════════════════════════
class _PlaywrightSession:
    """
    Playwright 브라우저 세션.

    ★ 항상 별도 스레드에서 Playwright를 실행합니다 ★
    - asyncio.get_running_loop() 감지에 의존하지 않음
    - Colab / Jupyter / FastAPI / 일반 Python 모두 동일하게 동작
    - sync_playwright()는 자체 이벤트 루프를 만들기 때문에
      메인 스레드의 asyncio 루프와 충돌 → 스레드 격리로 해결
    - render() 호출마다 독립 브라우저 인스턴스 생성·종료
      (세션 내 브라우저 재사용보다 안정성 우선)

    v4.0: Notion 렌더링 시 문서 높이가 안정될 때까지 반복 스크롤
          → lazy-load 블록까지 완전 로딩
    """

    def __init__(self, headless: bool = True, timeout_ms: int = 25000):
        self._headless   = headless
        self._timeout_ms = timeout_ms

    def __enter__(self): return self
    def __exit__(self, *_): pass

    @property
    def available(self) -> bool:
        return HAS_PLAYWRIGHT

    def _run_in_thread(self, fn, *args, **kwargs) -> str | None:
        """fn을 새 스레드에서 실행하고 결과 반환. timeout 초과 시 None."""
        result = [None]
        error  = [None]
        timeout_sec = self._timeout_ms / 1000 + 30

        def _target():
            try:
                result[0] = fn(*args, **kwargs)
            except Exception as e:
                error[0] = e

        t = threading.Thread(target=_target, daemon=True)
        t.start()
        t.join(timeout=timeout_sec)

        if not t.is_alive() and error[0]:
            print(f"    [Playwright] 스레드 에러: {type(error[0]).__name__}: {error[0]}",
                  flush=True)
        elif t.is_alive():
            print(f"    [Playwright] 타임아웃 ({timeout_sec}s 초과)", flush=True)

        return result[0]

    def _render_impl(self, url: str, notion: bool, wait: str) -> str | None:
        """실제 렌더링 로직. 반드시 메인 스레드가 아닌 곳에서 호출."""
        try:
            with sync_playwright() as p:
                browser = p.chromium.launch(headless=self._headless)
                try:
                    page = browser.new_page(
                        user_agent=(
                            "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                            "AppleWebKit/537.36 (KHTML, like Gecko) "
                            "Chrome/124.0 Safari/537.36"
                        )
                    )
                    page.set_extra_http_headers(
                        {"Accept-Language": "ko-KR,ko;q=0.9,en-US;q=0.8"}
                    )
                    if notion:
                        page.goto(url, wait_until="load", timeout=self._timeout_ms)
                        page.wait_for_timeout(2500)
                        # 높이 안정화까지 반복 스크롤 (lazy-load 완전 로딩)
                        prev_height = -1
                        for _ in range(8):
                            height = page.evaluate("document.body.scrollHeight")
                            if height == prev_height:
                                break
                            prev_height = height
                            page.evaluate("window.scrollTo(0, document.body.scrollHeight)")
                            page.wait_for_timeout(1200)
                        page.evaluate("window.scrollTo(0, 0)")
                        page.wait_for_timeout(500)
                    else:
                        try:
                            page.goto(url, wait_until=wait, timeout=self._timeout_ms)
                        except PWTimeoutError:
                            if wait == "networkidle":
                                print(f"    [Playwright] networkidle 타임아웃 → load 재시도",
                                      flush=True)
                            page.goto(url, wait_until="load", timeout=self._timeout_ms)
                            page.wait_for_timeout(2000)

                    html = page.content()
                    label = "Notion" if notion else "SPA"
                    print(f"    [Playwright] {label} 렌더링 완료: {len(html)} chars",
                          flush=True)
                    return html
                finally:
                    browser.close()
        except Exception as e:
            print(f"    [Playwright] 렌더링 실패 [{url}]: {type(e).__name__}: {e}",
                  flush=True)
            return None

    def render(self, url: str, wait: str = "networkidle") -> str | None:
        """SPA 페이지 렌더링. 항상 스레드에서 실행."""
        if not self.available:
            return None
        return self._run_in_thread(self._render_impl, url, False, wait)

    def render_notion(self, url: str) -> str | None:
        """Notion 페이지 전용 렌더링. 항상 스레드에서 실행."""
        if not self.available:
            return None
        return self._run_in_thread(self._render_impl, url, True, "load")


# ══════════════════════════════════════════════════════════════════
# 4-2. 딥 크롤러 (v4.0: 병렬 + 우선순위 큐 + sitemap + 중복 제거)
# ══════════════════════════════════════════════════════════════════
class _CrawlFrontier:
    """
    우선순위 기반 크롤 프론티어.
    - 페이지: (-priority, seq) 힙 → 커리어 관련 경로 우선 방문
    - PDF: 별도 큐 (항상 페이지보다 우선)
    - URL 정규화 키로 중복 등록 차단
    """

    def __init__(self, origin_url: str):
        self._origin = origin_url
        self._heap: list[tuple[int, int, str]] = []
        self._pdfs: deque[str] = deque()
        self._seen: set[str] = set()
        self._seq = 0

    def add_links(self, links: list[str]) -> None:
        for link in links:
            clean = link.split("#")[0].strip()
            if not clean or clean.startswith(("mailto:", "tel:", "javascript:", "data:")):
                continue
            norm = _normalize_url(clean)
            if not norm or norm in self._seen:
                continue
            if _should_skip_url(norm):
                continue
            self._seen.add(norm)
            if _is_pdf_url(norm):
                self._pdfs.append(norm)
            elif _same_origin(self._origin, norm):
                self._seq += 1
                heapq.heappush(self._heap, (-_url_priority(norm), self._seq, norm))

    def pop_batch(self, page_budget: int, pdf_budget: int, batch_size: int) -> list[tuple[str, bool]]:
        """(url, is_pdf) 목록 반환. PDF 우선, 이후 우선순위 높은 페이지."""
        batch: list[tuple[str, bool]] = []
        while self._pdfs and pdf_budget > 0 and len(batch) < batch_size:
            batch.append((self._pdfs.popleft(), True))
            pdf_budget -= 1
        while self._heap and page_budget > 0 and len(batch) < batch_size:
            _, _, url = heapq.heappop(self._heap)
            batch.append((url, False))
            page_budget -= 1
        return batch

    @property
    def has_pages(self) -> bool:
        return bool(self._heap)

    @property
    def has_pdfs(self) -> bool:
        return bool(self._pdfs)


def deep_crawl_site(start_url: str) -> str:
    """
    일반 웹사이트 딥 크롤러 (v4.0).
    - 병렬 fetch (기본 6 워커)로 크롤링 수 배 고속화
    - 커리어 관련 경로 우선 방문 (우선순위 큐)
    - sitemap.xml 자동 시딩
    - 콘텐츠 지문 기반 중복 페이지 제거 + 보일러플레이트(내비/푸터) 삭제
    - JSON-LD 구조화 데이터 추출
    - SPA 감지 시 Playwright로 완전 렌더링 (렌더링은 안정성 위해 순차 처리)
    동일 URL은 캐시에서 즉시 반환 (이중 크롤 방지).
    """
    cache_key = _normalize_url(start_url)
    if cache_key in _URL_CRAWL_CACHE:
        print(f"  [딥 크롤러] 캐시 히트 → 재크롤 생략: {start_url}", flush=True)
        return _URL_CRAWL_CACHE[cache_key]
    print(f"  [딥 크롤러 v4] 시작: {start_url}", flush=True)

    labels:  list[str] = []   # collected와 1:1 대응하는 헤더
    collected: list[str] = [] # 본문 텍스트 (보일러플레이트 제거 대상)
    extras:  list[str] = []   # PDF·JSON-LD 등 보일러플레이트 제거 제외 대상
    fingerprints: set[str] = set()
    pdf_count  = 0
    page_count = 0

    # ── 시작 페이지 ──────────────────────────────────────────────
    raw, ctype = _fetch(start_url)
    if raw is None:
        return ""
    if _is_pdf_url(start_url, raw):
        text = _parse_pdf_bytes(raw)
        result = f"[PDF: {start_url}]\n{text}" if text else ""
        _URL_CRAWL_CACHE[cache_key] = result
        return result

    main_text, main_links, main_html = _parse_html_bytes(raw, start_url, ctype)

    with _PlaywrightSession() as pw:
        # SPA 감지 시 Playwright로 재렌더링
        if _is_spa(main_html):
            if pw.available:
                print("    SPA detected → Playwright 렌더링 시작", flush=True)
                rendered = pw.render(start_url)
                if rendered:
                    main_text, main_links, main_html = _parse_html_bytes(
                        rendered.encode("utf-8", errors="replace"), start_url
                    )
            else:
                print("    SPA detected → Playwright 미설치, raw HTML 링크 추출로 폴백", flush=True)

        if main_text.strip():
            fingerprints.add(_content_fingerprint(main_text))
            labels.append(f"[Main: {start_url}]")
            collected.append(main_text)
            print(f"    Main page: {len(main_text)} chars", flush=True)

        json_ld = _extract_json_ld(main_html)
        if json_ld:
            extras.append(f"[구조화 데이터(JSON-LD): {start_url}]\n{json_ld}")

        # ── 프론티어 구성: 본문 링크 + raw HTML 링크 + sitemap ───
        frontier = _CrawlFrontier(start_url)
        frontier.add_links(main_links)
        frontier.add_links(_extract_links_from_raw_html(main_html, start_url))
        frontier.add_links(_fetch_sitemap_urls(start_url))

        print(f"    프론티어 초기화 완료 (병렬 워커: {_CRAWL_WORKERS})", flush=True)

        # ── 병렬 크롤 루프 ───────────────────────────────────────
        with ThreadPoolExecutor(max_workers=_CRAWL_WORKERS) as pool:
            while True:
                page_budget = _MAX_PAGES - page_count
                pdf_budget  = _MAX_PDFS - pdf_count
                if page_budget <= 0 and pdf_budget <= 0:
                    break
                if not frontier.has_pages and not frontier.has_pdfs:
                    break

                batch = frontier.pop_batch(page_budget, pdf_budget, _CRAWL_WORKERS)
                if not batch:
                    break

                futures = {pool.submit(_fetch, u): (u, is_pdf) for u, is_pdf in batch}
                spa_retry: list[str] = []

                for fut in as_completed(futures):
                    url, expected_pdf = futures[fut]
                    try:
                        data, sub_ctype = fut.result()
                    except Exception as e:
                        print(f"    WARNING fetch exception [{url}]: {e}", flush=True)
                        continue
                    if data is None:
                        continue

                    # PDF 처리 (확장자 또는 매직 바이트)
                    if _is_pdf_url(url, data):
                        if pdf_count >= _MAX_PDFS:
                            continue
                        text = _parse_pdf_bytes(data)
                        if text.strip():
                            extras.append(f"[PDF: {url}]\n{text}")
                            pdf_count += 1
                            print(f"    PDF collected ({pdf_count}/{_MAX_PDFS}): {url}", flush=True)
                        continue

                    if expected_pdf or page_count >= _MAX_PAGES:
                        continue

                    sub_text, sub_links, sub_html = _parse_html_bytes(data, url, sub_ctype)

                    # 서브페이지가 SPA면 렌더링 대상으로 표시 (렌더링은 순차)
                    if _is_spa(sub_html) and pw.available:
                        spa_retry.append(url)
                        continue

                    page_count += 1
                    fp = _content_fingerprint(sub_text)
                    if sub_text.strip() and fp not in fingerprints:
                        fingerprints.add(fp)
                        labels.append(f"[Page: {url}]")
                        collected.append(sub_text)
                        print(f"    Page {page_count}: {url} ({len(sub_text)} chars)", flush=True)
                    elif fp in fingerprints:
                        print(f"    중복 콘텐츠 스킵: {url}", flush=True)

                    json_ld = _extract_json_ld(sub_html)
                    if json_ld:
                        extras.append(f"[구조화 데이터(JSON-LD): {url}]\n{json_ld}")

                    frontier.add_links(sub_links)
                    frontier.add_links(_extract_links_from_raw_html(sub_html, url))

                # SPA 서브페이지는 순차 렌더링 (Playwright 안정성)
                for url in spa_retry:
                    if page_count >= _MAX_PAGES:
                        break
                    rendered = pw.render(url)
                    if not rendered:
                        continue
                    sub_text, sub_links, sub_html = _parse_html_bytes(
                        rendered.encode("utf-8", errors="replace"), url
                    )
                    page_count += 1
                    fp = _content_fingerprint(sub_text)
                    if sub_text.strip() and fp not in fingerprints:
                        fingerprints.add(fp)
                        labels.append(f"[Page(SPA): {url}]")
                        collected.append(sub_text)
                        print(f"    Page {page_count} (SPA): {url} ({len(sub_text)} chars)", flush=True)
                    frontier.add_links(sub_links)
                    frontier.add_links(_extract_links_from_raw_html(sub_html, url))

    # ── 보일러플레이트 제거 후 조립 ──────────────────────────────
    cleaned = _strip_boilerplate(collected)
    parts = [f"{label}\n{text}" for label, text in zip(labels, cleaned) if text.strip()]
    parts.extend(extras)

    result = "\n\n".join(parts)
    print(f"  [딥 크롤러 완료] 페이지 {page_count + 1}개 + PDF {pdf_count}개, {len(result)} chars", flush=True)
    _URL_CRAWL_CACHE[cache_key] = result
    return result


def read_file(path: str) -> str:
    """파일 경로를 받아 텍스트 내용 반환. PDF는 pypdf, DOCX는 python-docx로 파싱."""
    if not os.path.isfile(path):
        print(f"  ERROR file not found: {path}", flush=True)
        return ""
    lower = path.lower()
    if lower.endswith(".pdf"):
        with open(path, "rb") as f:
            return _parse_pdf_bytes(f.read())
    if lower.endswith(".docx"):
        if not HAS_DOCX:
            print("  WARNING: pip install python-docx  → DOCX 파싱 불가", flush=True)
            return ""
        try:
            doc = _docx.Document(path)
            paras = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        paras.append(" | ".join(cells))
            content = "\n".join(paras)
            print(f"  OK docx read ({len(content)} chars)", flush=True)
            return content
        except Exception as e:
            print(f"  ERROR reading docx: {e}", flush=True)
            return ""
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()
        print(f"  OK file read ({len(content)} chars)", flush=True)
        return content
    except Exception as e:
        print(f"  ERROR reading file: {e}", flush=True)
        return ""


# ══════════════════════════════════════════════════════════════════
# 5. 복합 입력 처리
# ══════════════════════════════════════════════════════════════════
_URL_PATTERN = re.compile(
    r'(https?://[^\s<>\[\]()\"\']+|www\.[^\s<>\[\]()\"\']+)',
    re.IGNORECASE
)
_FILE_PATH_PATTERN = re.compile(
    r'(?:^|(?<=\s))([/~]?(?:[\w.-]+/)*[\w.-]+\.(?:pdf|txt|md|json|csv|docx?|xlsx?|pptx?))(?=\s|$)',
    re.IGNORECASE | re.MULTILINE
)


def _extract_urls(text: str) -> list[str]:
    urls = []
    for match in _URL_PATTERN.finditer(text):
        url = match.group(1).strip().rstrip('.,;:!?')
        if url not in urls:
            urls.append(url)
    return urls


def _extract_file_paths(text: str) -> list[str]:
    paths = []
    for match in _FILE_PATH_PATTERN.finditer(text):
        path = match.group(1).strip()
        if path not in paths and os.path.isfile(path):
            paths.append(path)
    return paths


def _remove_patterns_from_text(text: str, patterns: list[str]) -> str:
    result = text
    for pattern in patterns:
        result = result.replace(pattern, " ")
    result = re.sub(r'\n{3,}', '\n\n', result)
    result = re.sub(r' {2,}', ' ', result)
    return result.strip()


def _resolve_mixed_input(source: str) -> str:
    source = source.strip()
    if not source:
        return ""

    collected_parts: list[str] = []

    urls = _extract_urls(source)
    if urls:
        print(f"\n  [복합 입력] {len(urls)}개 URL 감지됨", flush=True)
        for url in urls:
            full_url = url if url.startswith("http") else "https://" + url
            print(f"    → 크롤링: {full_url}", flush=True)
            if _is_notion_url(full_url):
                content = crawl_notion_page(full_url)
            else:
                content = deep_crawl_site(full_url)
            if content:
                collected_parts.append(f"[SOURCE_URL: {full_url}]\n{content}")

    file_paths = _extract_file_paths(source)
    if file_paths:
        print(f"\n  [복합 입력] {len(file_paths)}개 파일 감지됨", flush=True)
        for filepath in file_paths:
            content = read_file(filepath)
            if content:
                collected_parts.append(f"[SOURCE_FILE: {filepath}]\n{content}")

    remaining_text = _remove_patterns_from_text(source, urls + file_paths)
    if remaining_text and len(remaining_text.strip()) >= 10:
        print(f"\n  [복합 입력] 텍스트 {len(remaining_text)}자 감지됨", flush=True)
        collected_parts.append(f"[SOURCE_TEXT]\n{remaining_text}")

    if not collected_parts:
        print(f"\n  [텍스트 감지] {len(source)}자", flush=True)
        return source

    result = "\n\n" + "=" * 50 + "\n\n".join(collected_parts)
    print(f"\n  [복합 입력 완료] 총 {len(collected_parts)}개 소스, {len(result)} chars", flush=True)
    return result


def _resolve_input(source: str) -> str:
    """
    입력 유형 자동 판별 라우터 (v4.0 재작성).
    v3.0 버그 수정: 단일 URL 분기에서 크롤 결과를 반환하지 않고 버리던 문제 해결.
    URL/파일이 하나라도 감지되면 _resolve_mixed_input이 전부 처리합니다
    (크롤 결과는 URL 캐시로 이중 작업 없음).
    """
    source = source.strip()
    if not source:
        return ""

    # 입력 전체가 실존 파일 경로인 경우
    if os.path.isfile(source):
        print(f"\n  [파일 감지] 읽는 중: {source}", flush=True)
        return read_file(source)

    urls       = _extract_urls(source)
    file_paths = _extract_file_paths(source)

    if urls or file_paths:
        return _resolve_mixed_input(source)

    print(f"\n  [텍스트 감지] {len(source)}자", flush=True)
    return source


# ══════════════════════════════════════════════════════════════════
# 6. 중복 제거 & 병합
# ══════════════════════════════════════════════════════════════════
def generate_career_id(career: dict) -> str:
    """title + organization + period_start 조합으로 안정적인 ID 생성."""
    title        = (career.get("title") or "").strip().lower()
    organization = (career.get("organization") or "").strip().lower()
    period_start = (career.get("period_start") or "").strip()
    if not title:
        title = (career.get("description") or "")[:50].strip().lower()
    key = f"{title}|{organization}|{period_start}"
    return hashlib.md5(key.encode("utf-8")).hexdigest()[:12]


def deduplicate_careers(careers: list[dict]) -> list[dict]:
    """중복 경력 제거 후 각 항목에 'id' 부여. 동일 ID는 마지막 항목 유지."""
    seen: dict[str, dict] = {}
    for career in careers:
        c = dict(career)
        c["id"] = generate_career_id(career)
        seen[c["id"]] = c
    return list(seen.values())


def merge_careers(existing: list[dict], new_careers: list[dict]) -> list[dict]:
    """두 경력 리스트를 중복 없이 병합. 충돌 시 existing 우선."""
    seen: dict[str, dict] = {}
    for career in existing:
        c = dict(career)
        c["id"] = generate_career_id(career)
        seen[c["id"]] = c

    added = 0
    for career in new_careers:
        c = dict(career)
        c["id"] = generate_career_id(career)
        if c["id"] not in seen:
            seen[c["id"]] = c
            added += 1

    print(f"  [병합 결과] 기존 {len(existing)}개 + 신규 {added}개 → 총 {len(seen)}개", flush=True)
    return list(seen.values())


# ══════════════════════════════════════════════════════════════════
# 7. K-NN 파이프라인 (v4.0: 하이브리드 스코어 + 멀티 팩싯 + MMR)
# ══════════════════════════════════════════════════════════════════

def build_text_repr(career: dict) -> str:
    """
    임베딩용 텍스트 표현 생성.
    임베딩 모델이 앞쪽에 더 가중치를 두므로 중요도 순서대로 배치.
    title·role → description → achievements → raw_content(뒤, 노이즈 포함)
    """
    parts = []
    if career.get("title"):
        parts.append(f"제목: {career['title']}")
    if career.get("category"):
        parts.append(f"유형: {career['category']}")
    if career.get("role"):
        parts.append(f"역할: {career['role']}")
    if career.get("description"):
        parts.append(f"내용: {career['description']}")
    if career.get("achievements"):
        ach = career["achievements"]
        parts.append(f"성과: {', '.join(ach) if isinstance(ach, list) else ach}")
    if career.get("raw_content"):
        # raw_content는 뒤에 — 정보 풍부하지만 노이즈도 많음
        parts.append(career["raw_content"][:1500])
    return "\n".join(parts)


def _keyword_category(keyword: str) -> str:
    """키워드가 속한 카테고리 반환 (없으면 빈 문자열). 확장 폴백에 활용."""
    for category, kws in KEYWORD_CATEGORIES.items():
        if keyword in kws:
            return category
    return ""


# 키워드 확장 캐시 — analyze_keywords 재호출 시 LLM/임베딩 호출 생략
_KEYWORD_EXPANSION_CACHE: dict[str, dict] = {}


def expand_keyword(keyword: str) -> dict:
    """
    키워드를 멀티 팩싯(다관점)으로 확장 (v4.0).
    단어 하나는 임베딩하기에 너무 sparse하므로 LLM으로 4개 관점을 생성:
      - definition:       키워드의 정의 (1~2문장)
      - behaviors:        이 역량이 드러나는 구체적 행동 표현 (3~5개)
      - synonyms:         동의어·유사 표현 (3~5개)
      - evidence_phrases: 경력 기술서에 실제로 등장할 법한 성과 문장 (2~4개)

    각 팩싯을 개별 임베딩 후 센트로이드를 취하면 단일 문장 확장보다
    의미 공간을 넓게 커버하여 재현율(recall)이 향상됩니다.

    Returns:
        {"keyword", "definition", "behaviors", "synonyms",
         "evidence_phrases", "combined_text", "facet_texts"}
    실패 시 카테고리 힌트 기반 폴백.
    """
    if keyword in _KEYWORD_EXPANSION_CACHE:
        cached = _KEYWORD_EXPANSION_CACHE[keyword]
        # 임베딩까지 캐시된 경우 재사용 표시
        print(f"  [쿼리 확장] '{keyword}' 캐시 재사용", flush=True)
        return cached

    _init_client()
    expansion: dict | None = None
    try:
        def _do():
            return client.models.generate_content(
                model=_ANALYSIS_MODEL,
                contents=(
                    f"역량 키워드 '{keyword}'를 임베딩 검색에 사용할 수 있도록 확장하세요.\n"
                    f"다음 JSON 형식으로만 출력하세요:\n"
                    f'{{\n'
                    f'  "definition": "키워드 정의 1~2문장",\n'
                    f'  "behaviors": ["이 역량이 드러나는 구체적 행동 표현 3~5개"],\n'
                    f'  "synonyms": ["동의어/유사 표현 3~5개"],\n'
                    f'  "evidence_phrases": ["경력 기술서에 실제 등장할 법한 성과 문장 2~4개"]\n'
                    f'}}\n'
                    f"모두 한국어로 작성하세요."
                ),
                config=types.GenerateContentConfig(
                    temperature=0.1,
                    max_output_tokens=800,
                    response_mime_type="application/json",
                ),
            )
        resp = _call_with_retry(_do)
        expansion = json.loads(_clean_json_response(resp.text.strip()))
    except Exception as e:
        print(f"  [쿼리 확장 실패] '{keyword}': {e} → 카테고리 폴백", flush=True)

    if not isinstance(expansion, dict) or not expansion.get("definition"):
        category = _keyword_category(keyword)
        cat_hint = f" ({category} 관련 키워드)" if category else ""
        expansion = {
            "definition": f"{keyword}{cat_hint}: 경력과 활동에서 {keyword}이(가) 드러난 경험",
            "behaviors": [], "synonyms": [], "evidence_phrases": [],
        }

    definition = str(expansion.get("definition", keyword))
    behaviors  = [str(b) for b in expansion.get("behaviors", []) if b][:5]
    synonyms   = [str(s) for s in expansion.get("synonyms", []) if s][:5]
    evidences  = [str(p) for p in expansion.get("evidence_phrases", []) if p][:4]

    # 팩싯 텍스트 구성 (최대 3개 임베딩 → 센트로이드)
    facet_texts = [f"{keyword}. {definition}"]
    if behaviors:
        facet_texts.append(f"{keyword} 관련 행동: " + " / ".join(behaviors))
    facet_block = []
    if synonyms:
        facet_block.append("유사 표현: " + ", ".join(synonyms))
    if evidences:
        facet_block.append("성과 사례: " + " / ".join(evidences))
    if facet_block:
        facet_texts.append(f"{keyword} — " + " | ".join(facet_block))

    combined_text = "\n".join(
        [f"{keyword}: {definition}"]
        + ([f"행동: {', '.join(behaviors)}"] if behaviors else [])
        + ([f"동의어: {', '.join(synonyms)}"] if synonyms else [])
        + ([f"성과 표현: {' / '.join(evidences)}"] if evidences else [])
    )

    result = {
        "keyword": keyword,
        "definition": definition,
        "behaviors": behaviors,
        "synonyms": synonyms,
        "evidence_phrases": evidences,
        "combined_text": combined_text,
        "facet_texts": facet_texts,
    }
    _KEYWORD_EXPANSION_CACHE[keyword] = result
    print(f"  [쿼리 확장] '{keyword}' → 팩싯 {len(facet_texts)}개 / {len(combined_text)}자", flush=True)
    return result


def _keyword_embedding(keyword: str) -> tuple[Optional[list], str]:
    """
    키워드의 검색용 임베딩 생성 (v4.0).
    멀티 팩싯 각각을 RETRIEVAL_QUERY 태스크로 임베딩 후 센트로이드 반환.
    확장 캐시에 임베딩도 함께 저장하여 재호출 시 API 호출 0회.

    Returns:
        (centroid_embedding | None, combined_text)
    """
    expansion = expand_keyword(keyword)
    if expansion.get("_embedding"):
        return expansion["_embedding"], expansion["combined_text"]

    vectors = []
    for facet in expansion["facet_texts"]:
        vec = get_embedding(facet, task_type=_EMB_TASK_QUERY)
        if vec:
            vectors.append(vec)

    if not vectors:
        return None, expansion["combined_text"]

    dim = len(vectors[0])
    centroid = [sum(v[i] for v in vectors) / len(vectors) for i in range(dim)]
    centroid = _unit_normalize(centroid)

    expansion["_embedding"] = centroid  # 캐시에 임베딩 저장
    return centroid, expansion["combined_text"]


def _mmr_rerank(
    scored: list[tuple[dict, float]],
    k: int,
    lambda_: float = _KNN_MMR_LAMBDA,
) -> list[tuple[dict, float]]:
    """
    MMR(Maximal Marginal Relevance) 재랭킹.
    관련도가 높으면서도 이미 선택된 경험과 중복되지 않는 경험을 우선 선택
    → 같은 프로젝트의 파생 경험들이 상위를 독점하는 문제 방지.

    Args:
        scored: [(career, hybrid_score)] — 점수 내림차순 정렬 상태
        k:      선택할 개수
        lambda_: 1.0=관련도만 고려, 0.0=다양성만 고려
    """
    if len(scored) <= k:
        return scored

    selected = [scored[0]]
    remaining = list(scored[1:])

    while remaining and len(selected) < k:
        def _mmr_value(item):
            career, score = item
            emb = career.get("embedding")
            if not emb:
                return lambda_ * score
            max_sim = max(
                cosine_similarity(emb, s[0].get("embedding") or [])
                for s in selected
            )
            return lambda_ * score - (1 - lambda_) * max_sim

        best = max(remaining, key=_mmr_value)
        selected.append(best)
        remaining.remove(best)

    return selected


def knn_filter(
    query_embedding: list,
    careers: list[dict],
    k: int = _KNN_K,
    threshold: float = _KNN_SIM_THRESHOLD,
    query_text: str = "",
) -> list[tuple[dict, float]]:
    """
    하이브리드 K-NN 필터링 (v4.0).

    점수 산식:
      hybrid = _KNN_COSINE_WEIGHT × cosine(질의, 경험)
             + _KNN_LEXICAL_WEIGHT × lexical_containment(질의 어휘, 경험 텍스트)

    선별 로직 (3단계):
      1) 절대 컷: hybrid >= threshold 인 후보만 통과
      2) 상대 컷: hybrid >= (최고점 × _KNN_RELATIVE_CUT) — 상위권 대비
         현저히 낮은 후보 탈락 (전반적 유사도가 높은 풀에서 노이즈 차단)
      3) MMR 재랭킹으로 상위 k개 선택 — 중복 경험 독점 방지

    Returns:
        [(career_dict, hybrid_score), ...] — 점수 내림차순
    """
    scored: list[tuple[dict, float]] = []
    for c in careers:
        emb = c.get("embedding")
        if not emb:
            continue
        cos = cosine_similarity(query_embedding, emb)
        lex = 0.0
        if query_text:
            doc_text = c.get("_text_repr") or build_text_repr(c)
            lex = lexical_containment(query_text, doc_text)
        hybrid = _KNN_COSINE_WEIGHT * cos + _KNN_LEXICAL_WEIGHT * lex
        if hybrid >= threshold:
            scored.append((c, round(hybrid, 4)))

    if not scored:
        return []

    scored.sort(key=lambda x: x[1], reverse=True)

    # 상대 컷: 최고점 대비 현저히 낮은 후보 제거 (최소 1개는 유지)
    top_score = scored[0][1]
    relative_cut = top_score * _KNN_RELATIVE_CUT
    kept = [s for s in scored if s[1] >= relative_cut] or scored[:1]

    return _mmr_rerank(kept, k)


def _knn_select(
    keywords: list[str],
    careers: list[dict],
    k: int = _KNN_K,
    threshold: float = _KNN_SIM_THRESHOLD,
) -> dict[str, list[tuple[dict, float]]]:
    """
    키워드 목록별로 하이브리드 K-NN 필터링 수행.
    임베딩 생성 실패 시 어휘 유사도 단독 랭킹으로 폴백
    (v3.0의 '전체 경험 폴백'보다 정밀 — LLM 컨텍스트 낭비 방지).

    Returns:
        {keyword: [(career, hybrid_score), ...]}
    """
    result: dict[str, list[tuple[dict, float]]] = {}

    for kw in keywords:
        print(f"\n  [K-NN] '{kw}' 처리 중...", flush=True)
        kw_emb, combined_text = _keyword_embedding(kw)

        if kw_emb:
            top = knn_filter(kw_emb, careers, k=k, threshold=threshold,
                             query_text=combined_text)
            result[kw] = top
            sims = [f"{s:.3f}" for _, s in top]
            print(f"  [K-NN] '{kw}' → {len(top)}개 선별 | 하이브리드 점수: {sims}", flush=True)
        else:
            # 임베딩 실패 → 어휘 유사도 단독 랭킹 폴백
            print(f"  [K-NN] '{kw}' 임베딩 실패 → 어휘 유사도 폴백", flush=True)
            lex_scored = []
            for c in careers:
                doc_text = c.get("_text_repr") or build_text_repr(c)
                lex = lexical_containment(combined_text or kw, doc_text)
                lex_scored.append((c, round(lex, 4)))
            lex_scored.sort(key=lambda x: x[1], reverse=True)
            result[kw] = lex_scored[:max(k, _KNN_K)]

    return result


def explain_knn(
    keyword: str,
    careers: list[dict],
    top_n: int = 15,
) -> None:
    """
    K-NN 매칭 진단 유틸리티 (v4.0 신규).
    키워드에 대한 전체 경험의 코사인·어휘·하이브리드 점수와
    임계값 통과 여부를 표로 출력 → 임계값 튜닝에 활용.
    """
    print("=" * 72)
    print(f"  K-NN 매칭 진단: '{keyword}'")
    print("=" * 72)

    kw_emb, combined_text = _keyword_embedding(keyword)
    if not kw_emb:
        print("  ❌ 키워드 임베딩 생성 실패 — test_embedding()으로 진단하세요.")
        return

    rows = []
    for c in careers:
        emb = c.get("embedding")
        title = (c.get("title") or "제목 없음")[:28]
        if not emb:
            rows.append((title, None, None, None))
            continue
        cos = cosine_similarity(kw_emb, emb)
        doc_text = c.get("_text_repr") or build_text_repr(c)
        lex = lexical_containment(combined_text, doc_text)
        hybrid = _KNN_COSINE_WEIGHT * cos + _KNN_LEXICAL_WEIGHT * lex
        rows.append((title, cos, lex, hybrid))

    rows.sort(key=lambda r: (r[3] is not None, r[3] or 0), reverse=True)

    print(f"\n  절대 임계값: {_KNN_SIM_THRESHOLD} | 상대 컷: top×{_KNN_RELATIVE_CUT}"
          f" | 가중치: cos {_KNN_COSINE_WEIGHT} / lex {_KNN_LEXICAL_WEIGHT}\n")
    print(f"  {'경험':<30} {'cosine':>8} {'lexical':>8} {'hybrid':>8}  판정")
    print("  " + "-" * 66)
    top_score = max((r[3] for r in rows if r[3] is not None), default=0)
    for title, cos, lex, hybrid in rows[:top_n]:
        if hybrid is None:
            print(f"  {title:<30} {'—':>8} {'—':>8} {'—':>8}  임베딩 없음")
            continue
        passed_abs = hybrid >= _KNN_SIM_THRESHOLD
        passed_rel = hybrid >= top_score * _KNN_RELATIVE_CUT
        verdict = "✅ 통과" if (passed_abs and passed_rel) else \
                  ("△ 상대컷 탈락" if passed_abs else "✗ 임계값 미달")
        print(f"  {title:<30} {cos:>8.3f} {lex:>8.3f} {hybrid:>8.3f}  {verdict}")
    print("=" * 72)


def _build_careers_text_knn(
    knn_result: dict[str, list[tuple[dict, float]]],
    total_count: int,
) -> tuple[str, str]:
    """
    K-NN 모드용 careers_text 및 coverage_note 생성.
    raw_content는 컨텍스트가 이미 줄었으므로 전체 포함.

    Returns:
        (careers_text, coverage_note)
    """
    careers_text   = ""
    coverage_lines = []

    for kw, top_list in knn_result.items():
        related = len(top_list)
        pct     = round(related / total_count * 100, 1) if total_count > 0 else 0.0
        coverage_lines.append(f"- {kw}: {related}개 / 전체 {total_count}개 = {pct}%")

        careers_text += (
            f"\n\n{'='*60}\n"
            f"[키워드: '{kw}' — K-NN 선별 {related}개 / 전체 {total_count}개]\n"
            f"{'='*60}\n"
        )

        if not top_list:
            careers_text += "  ※ 유사도 기준을 충족하는 경험이 없습니다.\n"
            continue

        for rank, (career, sim) in enumerate(top_list, 1):
            period = (
                f"{career.get('period_start', '')} ~ {career.get('period_end', '')}"
                if career.get("period_start") or career.get("period_end")
                else "시기 미상"
            )
            ach = career.get("achievements", [])

            careers_text += (
                f"\n--- [{rank}위] 유사도: {sim:.3f} ---\n"
                f"제목: {career.get('title', '')}\n"
                f"유형: {career.get('category', '')}\n"
                f"기관: {career.get('organization', '')}\n"
                f"역할: {career.get('role', '')}\n"
                f"기간: {period}\n"
                f"설명: {career.get('description', '')}\n"
            )
            if ach:
                careers_text += f"성과: {', '.join(ach) if isinstance(ach, list) else ach}\n"
            if career.get("raw_content"):
                # K-NN 모드: 컨텍스트가 줄었으므로 raw_content 전체 포함
                careers_text += f"[원문]\n{career['raw_content']}\n"

    coverage_note = "\n".join(coverage_lines)
    return careers_text, coverage_note


def _build_careers_text_llm(careers: list[dict]) -> str:
    """
    LLM 직접 분석 모드용 careers_text 생성.
    raw_content 2,000자 절삭 (컨텍스트 보호).
    """
    careers_text = ""
    for i, career in enumerate(careers, 1):
        period = (
            f"{career.get('period_start', '')} ~ {career.get('period_end', '')}"
            if career.get("period_start") or career.get("period_end")
            else "시기 미상"
        )
        ach = career.get("achievements", [])

        careers_text += f"\n\n[경력 {i}] ID: {career.get('id', f'career_{i}')}\n"
        careers_text += f"제목: {career.get('title', '제목 없음')}\n"
        careers_text += f"카테고리: {career.get('category', '')}\n"
        careers_text += f"기관: {career.get('organization', '')}\n"
        careers_text += f"역할: {career.get('role', '')}\n"
        careers_text += f"기간: {period}\n"
        careers_text += f"설명: {career.get('description', '')}\n"
        if ach:
            careers_text += f"성과: {', '.join(ach) if isinstance(ach, list) else ach}\n"
        if career.get("raw_content"):
            careers_text += f"[원문]\n{career['raw_content'][:2000]}\n"

    return careers_text


# ══════════════════════════════════════════════════════════════════
# 8. 프롬프트
# ══════════════════════════════════════════════════════════════════
def _career_parsing_prompt() -> str:
    return """당신은 경력 데이터 파싱 전문가이며 이브누입니다.

[임무]
사용자의 경력/활동 데이터에서 개별 경력 항목을 추출합니다.

[★ HALLUCINATION 절대 금지 ★]
1. 입력 데이터에서 확인된 내용만 추출합니다.
2. 없는 정보는 null 또는 빈 배열로 처리합니다.
3. 날짜가 불명확하면 null. 절대 추측하지 마세요.
4. 기관명/회사명은 원문 그대로 사용합니다.
5. 성과/수치는 원문에 명시된 것만 사용합니다.
6. 원문에 없는 역할/직책을 창작하지 않습니다.

[추출 항목]
- title:        경력/활동 제목 (간결하게, 30자 이내)
- category:     인턴십|프로젝트|동아리|봉사|수상|자격증|연구|창업|공모전|대외활동|학회|교육|기타
- organization: 소속 기관/회사 (원문 그대로, 없으면 null)
- role:         역할/직책 (원문 그대로, 없으면 null)
- period_start: YYYY-MM 또는 YYYY. 불명확하면 null
- period_end:   YYYY-MM 또는 YYYY 또는 "진행중". 불명확하면 null
- description:  활동 내용 요약 (2~3문장, 원문 기반)
- achievements: 성과/결과 목록 (배열, 원문에 명시된 것만)
- raw_content:  해당 경력의 원문 텍스트 (최대 500자로 절삭, 큰따옴표/역슬래시 제거)

[날짜 파싱 규칙]
- "2023년 1월" → "2023-01"
- "2023.01" / "2023/01" / "Jan 2023" → "2023-01"
- "2023" → "2023"
- "현재" / "Present" / "진행중" → "진행중"
- 날짜가 없거나 애매하면 → null

[출력 형식 - 순수 JSON만]
{
  "status": "success",
  "careers": [
    {
      "title": "...",
      "category": "...",
      "organization": "...",
      "role": "...",
      "period_start": "...",
      "period_end": "...",
      "description": "...",
      "achievements": ["..."],
      "raw_content": "..."
    }
  ]
}"""


def _keyword_analysis_prompt(keywords: list[str], target_scenario: str = "") -> str:
    keywords_str = ", ".join(keywords)
    target_hint  = f"\n목표 시나리오/직무: {target_scenario}" if target_scenario else ""

    return f"""당신은 키워드 기반 경력 분석 전문가입니다.

[핵심 임무]
사용자가 선택한 키워드를 기준으로, 제공된 경험 풀에서 해당 키워드를 보여줄 수 있는 근거를 찾고 분석합니다.

[분석 대상 키워드]
{keywords_str}{target_hint}

[★ HALLUCINATION 절대 금지 ★]
1. 원문에 없는 내용을 창작/추측하지 않습니다.
2. source_quote는 반드시 원문에서 그대로 인용합니다 (최소 15자).
3. 원문에 없는 성과/수치/결과를 만들어내지 않습니다.
4. 근거가 부족하면 "근거 부족"으로 표시하고 과장하지 않습니다.
5. 확신도(confidence)는 근거 강도에 따라 정직하게 설정합니다.

[분석 원칙]
1. 키워드 단어 등장만으로 부합 판단 금지 (행동/결과/학습과 연결 필요)
2. 모든 평가/추천에 최소 1개 이상 근거 필수
3. 근거 강도에 따라 high/medium/low를 정직하게 부여할 것 — 과대 또는 과소 평가 금지
4. 날짜가 null인 경우 "시기 미상"으로 표시
5. 한국어를 기본 출력 언어로 설정
6. "데이터 분석"과 같이 특정 활동에 대해 분석할 때는 과한 추론을 금지하고,
   실제로 해당 분야 서비스·프로그램을 만든 경우를 최우선 기준으로 판별

[★ 억지 매칭 금지 — 포함보다 제외가 원칙 ★]
1. 키워드와 경험의 연결에 2단계 이상의 추론이 필요하면 그 경험은 제외합니다.
   (예: "봉사활동을 했다 → 사람을 만났다 → 커뮤니케이션 역량" 같은 연쇄 추론 금지)
2. 키워드당 경험 개수를 채우기 위해 무리하게 포함하지 않습니다.
   해당 키워드에 진짜 부합하는 경험이 1개뿐이면 1개만 반환하는 것이 정답입니다.
3. relevance=low는 "원문에서 직접 인용 가능한 근거가 최소 1개" 있을 때만 허용합니다.
   인용할 문장조차 없는 경험은 low로도 포함하지 말고 제외하십시오.
4. K-NN이 선별해 전달한 경험이라도 위 기준에 미달하면 제외할 수 있습니다.
   (K-NN은 의미 유사도만 봅니다 — 근거 실재 여부의 최종 판단은 당신의 몫입니다)
5. 제외한 경험이 있으면 B_selection_criteria.summary에 "N개 경험을 근거 부족으로
   제외함"을 명시하여 투명하게 남기십시오.

[관련도(relevance) 3단계 기준]
- high  : 키워드를 핵심 역량으로 직접 발휘한 경험. 구체적 행동+성과 모두 확인됨
- medium: 키워드와 관련 있으나 보조적 역할이거나 성과 근거가 일부 부족한 경험
- low   : 키워드와 연관성이 있지만 간접적이거나 근거가 매우 빈약한 경험.
          포함하되 "참고용"으로 명확히 표시하고 개선 방향을 함께 제시할 것

[confidence 3단계 기준]
- high  : 원문에서 직접 인용 가능한 근거 2개 이상 확인
- medium: 원문 근거 1개 또는 맥락 추론 가능
- low   : 원문 근거 불충분. 포함하되 confidence_reason에 한계를 명시할 것

[K-NN 모드 추가 지시 — user_prompt에 K-NN 데이터가 있는 경우]
- 제공된 경험은 하이브리드 유사도 기반으로 이미 선별된 경험입니다.
- 목록에 없는 경험을 추가로 찾아 포함하지 않습니다.
- 단, 선별된 경험이라도 [억지 매칭 금지] 기준에 미달하면 제외해야 합니다.
  선별 목록은 "후보"이지 "확정"이 아닙니다.
- total_count는 user_prompt의 값을 그대로 사용합니다.
- related_count와 coverage_percent는 user_prompt의 실측값을 기준으로 하되,
  당신이 근거 부족으로 제외한 경험이 있으면 그만큼 차감하여 반영합니다.

[스토리라인 작성 원칙 — E_storylines]
★ 가장 중요한 섹션. 단순 나열이 아닌 "하나의 이야기"를 만들어야 합니다.

━━━ 시간 순서 원칙 (필수) ━━━
1. period_start가 있는 경험은 반드시 시간 오름차순(오래된 → 최근)으로 배열
2. period_start가 null인 경험이 포함된 경우:
   - timeline_status: "일부_불명확" 또는 "대부분_불명확" 설정
   - timeline_note: 어떤 경험의 시점이 불명확한지 목록으로 명시
   - narrative에서 해당 경험 앞에 "(시점 불명확)" 명시
   - 날짜 없는 경험끼리의 상대 순서는 강제하지 않고 맥락상 자연스럽게 배치
   - chronological_sequence의 is_dated=false인 항목은 추정 위치 표시
3. 모든 경험에 날짜가 있으면 timeline_status: "시간순_확인됨"
4. narrative는 반드시 chronological_sequence 순서대로 작성

━━━ 서사 품질 원칙 (필수) ━━━
금지 패턴 — 절대 사용 금지:
  "A 경험에서 X를 했다. B 경험에서 Y를 했다."
  → 이력서 나열과 동일. 금지.

권장 패턴:
  "A에서 처음 X에 도전했지만 한계를 느꼈고, 그 부족함이 B로 이어지는 계기가 됐다.
   B에서 Y를 통해 방향을 잡았고, 그 경험이 C에서 Z를 가능하게 했다."
  → 인과·심화·반성 관계가 명확히 드러나야 함.

필수 작성 항목:
  - narrative: 시간 순서 기반 단일 서사 문단. 최소 150자.
               날짜 불명확 경험 앞에는 "(시점 불명확)" 표시.
               마지막 문장은 현재 역량/지향점으로 연결.
  - turning_points: 인식·행동이 실질적으로 바뀐 순간. 구체적 촉발 사건 명시.
  - connective_logic: 경험 A→B 연결 고리를 인과/심화/반성/확장으로 분류.

[개선 가이드 작성 원칙 — F_improvement_guide]
★ 나열식 제안 금지. 사용자가 "그래서 다음에 뭘 해야 하는가"를 즉시 알 수 있는
  명확한 방향성을 제시해야 합니다.

1. overall_direction (필수, 가장 먼저 작성):
   - current_profile_summary: 현재 경험 풀의 강점과 공백을 2~3문장으로 정직하게 진단.
     실제 분석 결과(C_coverage, D의 relevance 분포)에 근거할 것.
   - short_term: 앞으로 3개월 내 실행할 가장 효과 큰 액션 1~2개.
     "~을 하세요"가 아니라 "무엇을, 왜, 어떤 순서로"까지 구체적으로.
   - mid_term: 1년 내 보강 방향. 목표 시나리오(target)가 있으면 반드시 그에 정렬.
   - priority_keyword: 분석 키워드 중 가장 시급히 보강해야 할 키워드 1개와 이유.
     (coverage가 낮거나 high 근거가 없는 키워드가 우선 후보)
2. 모든 개선 항목(information_enhancement, experience_expansion)에
   priority("높음"|"중간"|"낮음")를 부여하고, 높음부터 정렬합니다.
3. 실행 불가능한 추상 제안 금지 ("리더십을 더 기르세요" ✗ /
   "현재 동아리 운영 경험에 정량 성과 수치를 추가 기록하세요" ✓).
4. 진단은 근거 기반으로: 각 제안이 어떤 분석 결과에서 나왔는지 reason에 연결.

[출력 형식 - A~F 순서 고정, 순수 JSON만]
{{
  "status": "success",
  "analysis_date": "YYYY-MM-DD",
  "analysis_mode": "knn|llm",
  "keywords": ["키워드1", "키워드2"],
  "target_scenario": "목표 시나리오",

  "A_keyword_definitions": [
    {{
      "keyword": "키워드명",
      "definition": "키워드 재정의 (1~2문장)",
      "synonyms": ["동의어", "유사표현"],
      "compliance_criteria": [
        {{
          "id": 1,
          "criterion": "부합 기준",
          "signal_description": "무엇을 보면 부합이라 보는지"
        }}
      ]
    }}
  ],

  "B_selection_criteria": {{
    "summary": "AI가 어떤 규칙으로 경험을 골랐는지 3~5줄 요약",
    "criteria": ["기준 매칭", "근거 강도", "반복성", "맥락 적합성"]
  }},

  "C_coverage": [
    {{
      "keyword": "키워드명",
      "related_count": 0,
      "total_count": 0,
      "coverage_percent": 0.0,
      "high_count": 0,
      "medium_count": 0,
      "low_count": 0
    }}
  ],

  "D_matched_experiences": [
    {{
      "keyword": "키워드명",
      "experiences": [
        {{
          "career_title": "경력 제목",
          "organization": "소속 기관",
          "period": "YYYY-MM ~ YYYY-MM 또는 시기 미상",
          "relevance": "high|medium|low",
          "relevance_summary": "이 경험이 키워드와 어느 정도 관련되는지 한 줄 요약",
          "evidence": [
            {{
              "type": "사건|행동|성과|학습|증빙",
              "content": "이 경험이 키워드와 관련되는 이유",
              "source_quote": "원문 그대로 인용 (최소 15자)"
            }}
          ],
          "matched_criteria": [1, 3],
          "confidence": "high|medium|low",
          "confidence_reason": "확신도 판단 이유 (모든 경우에 작성)"
        }}
      ]
    }}
  ],

  "E_storylines": [
    {{
      "keyword": "관련 키워드",
      "storyline_title": "스토리라인 제목 (한 줄, 이야기의 핵심 메시지)",
      "tagline": "면접에서 바로 쓸 수 있는 한 줄 자기소개 형식 요약 (30자 내외)",
      "timeline_status": "시간순_확인됨 | 일부_불명확 | 대부분_불명확",
      "timeline_note": "null 또는 시점 불명확 경험 목록 및 설명. 모두 날짜가 있으면 null.",
      "chronological_sequence": [
        {{
          "order": 1,
          "experience": "경험명",
          "period": "YYYY-MM 또는 YYYY 또는 시기_미상",
          "is_dated": true
        }}
      ],
      "narrative": "시간 순서 기반 서사형 단락. 날짜 불명확 경험 앞에 (시점 불명확) 표시. 최소 150자. 마지막 문장은 현재 역량/지향점으로 연결.",
      "turning_points": [
        {{
          "experience": "전환점이 발생한 경험명",
          "period": "해당 경험의 시점 (없으면 시기_미상)",
          "trigger": "변화를 촉발한 구체적 사건, 피드백, 결과 (원문 근거 기반)",
          "what_changed": "이 경험 전후로 무엇이 달라졌는지 (인식/행동/접근 방식)"
        }}
      ],
      "connective_logic": [
        {{
          "from_experience": "앞 경험명",
          "to_experience": "뒤 경험명",
          "relation_type": "인과|심화|반성|확장",
          "connection": "왜 앞 경험이 뒤 경험으로 이어졌는지 한 문장 설명",
          "temporal_note": "null 또는 시점 불명확시 '시점 미확인 연결' 명시"
        }}
      ],
      "structure": {{
        "start": "시작점: 이 역량에 처음 관심/필요를 느낀 계기 (시점 포함, 없으면 시기 불명확 명시)",
        "development": "전개: 시간 순서대로 경험들을 통해 역량을 쌓은 과정",
        "evidence": "증거: 역량을 가장 잘 보여주는 구체적 성과/산출물/피드백",
        "growth": "성장: 경험들을 거치며 달라진 점, 깨달은 것",
        "destination": "현재: 지금 이 역량이 어떤 수준이며 어떤 방향으로 발전 중인지"
      }},
      "used_experiences": {{
        "core": ["핵심 경력 — 스토리의 중심이 되는 경험"],
        "supporting": ["보조 경력 — 맥락을 보완하는 경험"]
      }},
      "key_quotes": [
        {{
          "career_title": "경력 제목",
          "quote": "원문 인용 (이야기의 흐름을 뒷받침하는 문장)"
        }}
      ]
    }}
  ],

  "F_improvement_guide": {{
    "overall_direction": {{
      "current_profile_summary": "현재 경험 풀의 강점·공백 진단 (2~3문장, 분석 결과 근거)",
      "short_term": "단기(3개월 내) 실행 방향 — 가장 효과 큰 액션 1~2개, 순서 포함",
      "mid_term": "중기(1년 내) 보강 방향 — 목표 시나리오에 정렬",
      "priority_keyword": "최우선 보강 키워드",
      "priority_reason": "그 키워드가 최우선인 이유 (coverage/근거 강도 기반)"
    }},
    "information_enhancement": [
      {{
        "target": "보강 대상 경력 제목",
        "missing": "부족한 정보",
        "how_to_add": "추가 방법",
        "reason": "이유 (어떤 분석 결과에서 나온 제안인지 연결)",
        "priority": "높음|중간|낮음"
      }}
    ],
    "experience_expansion": [
      {{
        "gap_description": "부족한 부분",
        "suggested_experience_type": "추천 경험 유형",
        "why_helpful": "이유",
        "examples": ["예시"],
        "priority": "높음|중간|낮음"
      }}
    ],
    "keyword_specific_recommendations": [
      {{
        "keyword": "키워드명",
        "recommendations": [
          {{"type": "확장|보완", "title": "추천 활동명", "expected_effect": "기대 효과"}}
        ]
      }}
    ]
  }}
}}"""


# ══════════════════════════════════════════════════════════════════
# 9. 공개 함수
# ══════════════════════════════════════════════════════════════════

def parse_careers(input_data: str) -> list[dict]:
    """
    URL / 파일경로 / 텍스트 / 복합 입력을 받아 파싱된 경력 리스트를 반환합니다.
    v4.0: 임베딩(RETRIEVAL_DOCUMENT)과 어휘 매칭용 텍스트 표현(_text_repr)을
          사전 생성하여 각 경력 dict에 캐시합니다.

    Args:
        input_data: URL, 파일경로, 텍스트, 또는 이들의 조합

    Returns:
        경력 dict 리스트. 각 항목에 'id', 'embedding', '_text_repr' 필드 자동 부여.
    """
    print("=" * 60)
    print("  경력 파싱 (v4.1 - Precision K-NN + Deep Crawler Edition)")
    print("=" * 60)

    raw_content = _resolve_input(input_data)
    if not raw_content or len(raw_content.strip()) < 10:
        print("[ERROR] 입력 데이터가 너무 짧습니다.")
        return []

    print("\n  [파싱 중...]", flush=True)
    sanitized = _sanitize_text(raw_content)
    result = _call_model(
        _career_parsing_prompt(),
        f"아래 텍스트에서 모든 경력/활동 항목을 추출하세요.\n\n[입력 데이터]\n{sanitized[:30000]}"
    )

    if result.get("status") == "error":
        print(f"  [ERROR] {result.get('message')}", flush=True)
        return []

    raw_careers = result.get("careers", [])
    careers     = deduplicate_careers(raw_careers)

    # ── 임베딩 사전 생성 (K-NN 파이프라인용) ─────────────────────
    print(f"\n  [임베딩 생성] {len(careers)}개 경력 처리 중...", flush=True)
    success_count = 0
    for c in careers:
        text_repr        = build_text_repr(c)
        c["_text_repr"]  = text_repr                      # 어휘 매칭용 캐시
        c["embedding"]   = get_embedding(text_repr, task_type=_EMB_TASK_DOC)
        if c["embedding"]:
            success_count += 1
    print(f"  [임베딩 완료] {success_count}/{len(careers)}개 성공", flush=True)
    # ─────────────────────────────────────────────────────────────

    print(f"\n[완료] {len(careers)}개 경력 반환 (원본 {len(raw_careers)}개 → 중복 제거 후 {len(careers)}개)")
    for i, c in enumerate(careers):
        has_emb = "✓" if c.get("embedding") else "✗"
        print(f"  {i+1}. [{c.get('id')}] {c.get('title', '제목 없음')} (임베딩:{has_emb})")

    return careers


def analyze_keywords(
    keywords: list[str],
    careers:  list[dict],
    target:   str  = "",
    pretty:   bool = True,
) -> dict:
    """
    키워드 리스트와 경력 리스트를 입력받아 분석 결과를 반환합니다.
    경험 수에 따라 자동으로 분석 모드를 선택합니다.
      - 경험 < _KNN_THRESHOLD_COUNT : LLM 직접 분석
      - 경험 >= _KNN_THRESHOLD_COUNT: K-NN 사전 필터링 → LLM 근거 추출
    v4.0: K-NN 필터링이 하이브리드 스코어(코사인+어휘) + 상대 컷 + MMR로 동작.

    Args:
        keywords: 분석할 키워드 목록 (1개 이상, 개수 제한 없음)
        careers:  parse_careers() 또는 merge_careers()의 반환값
        target:   목표 시나리오/직무 (선택)
        pretty:   True면 들여쓰기된 JSON 출력

    Returns:
        분석 결과 dict
    """
    print("=" * 60)
    print("  키워드 기반 분석 (v4.1)")
    print("=" * 60)

    # ── 입력 검증 ────────────────────────────────────────────────
    if isinstance(keywords, str):
        keywords = [k.strip() for k in keywords.split(",") if k.strip()]
    if not keywords:
        print("[ERROR] 최소 1개 이상의 키워드가 필요합니다.")
        return {"status": "error", "message": "키워드가 필요합니다."}
    # 키워드 수 제한 없음 — 개수가 많을수록 분석 시간과 API 비용 증가
    if not careers:
        print("[ERROR] careers가 비어 있습니다. parse_careers()로 먼저 경력을 파싱하세요.")
        return {"status": "error", "message": "경력 데이터 없음"}
    # ─────────────────────────────────────────────────────────────

    today_str  = datetime.now().strftime("%Y-%m-%d")
    total      = len(careers)
    use_knn    = total >= _KNN_THRESHOLD_COUNT

    print(f"\n분석 대상: {total}개 경력 / 키워드: {', '.join(keywords)}")
    if target:
        print(f"목표 시나리오: {target}")

    # ── 모드 분기 ────────────────────────────────────────────────
    if use_knn:
        print(f"\n  [모드] K-NN 파이프라인 (경험 {total}개 ≥ {_KNN_THRESHOLD_COUNT})", flush=True)
        knn_result   = _knn_select(keywords, careers)
        careers_text, coverage_note = _build_careers_text_knn(knn_result, total)
        analysis_mode = "knn"

        user_prompt = f"""아래 경력 풀을 분석하여 키워드 기반 분석 결과를 생성하세요.

[오늘 날짜]
{today_str}

[분석 모드]
K-NN 사전 필터링 완료. 이하 경험은 하이브리드 유사도(임베딩 코사인 + 어휘 일치)
기반으로 선별된 경험입니다.
LLM은 선별된 경험에 대한 근거 추출과 스토리라인 구성만 담당합니다.
추가로 관련 경험을 찾거나 포함 여부를 재판단하지 마세요.

[분석 대상 키워드]
{', '.join(keywords)}

[목표 시나리오/직무]
{target if target else '(없음)'}

[coverage_percent 실측값 — C_coverage에 이 수치를 그대로 사용할 것]
{coverage_note}

[K-NN 선별 경험 목록 (키워드별 · 유사도 내림차순)]
{careers_text}

[중요 지시]
- analysis_date는 반드시 "{today_str}"로 설정하세요.
- analysis_mode는 반드시 "knn"으로 설정하세요.
- source_quote는 원문에서 정확히 인용하세요 (최소 15자).
- C_coverage의 related_count, total_count, coverage_percent는
  위 [coverage_percent 실측값]의 수치를 그대로 사용하고 임의로 수정하지 마세요.
"""
    else:
        print(f"\n  [모드] LLM 직접 분석 (경험 {total}개 < {_KNN_THRESHOLD_COUNT}, K-NN 생략)", flush=True)
        careers_text  = _build_careers_text_llm(careers)
        analysis_mode = "llm"

        user_prompt = f"""아래 경력 풀을 분석하여 키워드 기반 분석 결과를 생성하세요.

[오늘 날짜]
{today_str}

[분석 모드]
LLM 직접 분석 (경험 수 {total}개로 K-NN 생략).

[분석 대상 키워드]
{', '.join(keywords)}

[목표 시나리오/직무]
{target if target else '(없음)'}

[전체 경력 풀 - {total}개]
{careers_text}

[중요 지시]
- analysis_date는 반드시 "{today_str}"로 설정하세요.
- analysis_mode는 반드시 "llm"으로 설정하세요.
- source_quote는 원문에서 정확히 인용하세요 (최소 15자).
- 원문에 없는 내용은 절대 인용하지 마세요.
"""
    # ─────────────────────────────────────────────────────────────

    result = _call_model(_keyword_analysis_prompt(keywords, target), user_prompt)

    if result.get("status") == "success":
        result["analysis_date"] = today_str
        result["analysis_mode"] = analysis_mode
        result = _enrich_and_tag_result(result)
        print(f"\n  [분석 완료] 모드: {analysis_mode.upper()}", flush=True)
    else:
        print(f"  [분석 실패] {result.get('message', 'Unknown error')}", flush=True)

    print("\n" + "=" * 60)
    print("  분석 결과 (JSON)")
    print("=" * 60)
    print(json.dumps(result, ensure_ascii=False, indent=2 if pretty else None))

    return result


def _enrich_and_tag_result(result: dict) -> dict:
    """
    분석 결과를 후처리합니다.
    - low 경험을 제거하지 않고 is_reference_only=True 태그를 부여
    - C_coverage에 high/medium/low 건수 집계 추가
    - 경험을 relevance 내림차순으로 정렬 (high → medium → low)
    """
    _order = {"high": 0, "medium": 1, "low": 2}
    matched  = result.get("D_matched_experiences", [])
    coverage = result.get("C_coverage", [])

    coverage_by_kw: dict[str, dict] = {c.get("keyword", ""): c for c in coverage}

    for kw_block in matched:
        kw   = kw_block.get("keyword", "")
        exps = kw_block.get("experiences", [])

        high_n = medium_n = low_n = 0
        for exp in exps:
            rel = exp.get("relevance", "medium")
            if rel == "high":
                high_n += 1
                exp["is_reference_only"] = False
            elif rel == "medium":
                medium_n += 1
                exp["is_reference_only"] = False
            else:
                low_n += 1
                exp["is_reference_only"] = True   # 참고용 태그

        # relevance 내림차순 정렬
        kw_block["experiences"] = sorted(
            exps, key=lambda e: _order.get(e.get("relevance", "medium"), 1)
        )

        # C_coverage 집계 업데이트
        if kw in coverage_by_kw:
            c      = coverage_by_kw[kw]
            total  = c.get("total_count", 0)
            related = high_n + medium_n + low_n
            c["related_count"]    = related
            c["high_count"]       = high_n
            c["medium_count"]     = medium_n
            c["low_count"]        = low_n
            c["coverage_percent"] = round(related / total * 100, 1) if total else 0.0

        if low_n:
            print(f"  [low 태깅] '{kw}': {low_n}개 경험 → is_reference_only=True", flush=True)

    return result


def show_keywords():
    """추천 키워드 목록을 출력합니다."""
    print("=" * 60)
    print("  추천 키워드 목록")
    print("=" * 60)
    for category, kws in KEYWORD_CATEGORIES.items():
        print(f"\n[{category}]")
        print(f"  {', '.join(kws)}")


def print_result_summary(result: dict) -> None:
    """
    분석 결과를 섹션별로 가독성 있게 출력합니다.
    - 3단계 관련도(high/medium/low) 아이콘으로 구분
    - 스토리라인 narrative + 전환점 + 연결 고리 전문 출력
    - low 경험은 [참고용] 표시
    """
    if result.get("status") != "success":
        print(f"[오류] {result.get('message', '분석 실패')}")
        return

    mode = result.get("analysis_mode", "").upper()
    date = result.get("analysis_date", "")
    kws  = ", ".join(result.get("keywords", []))
    W    = 64

    def sep(char="━"): print(char * W)
    def section(title): print(f"\n  {'─'*4} {title} {'─'*(W-8-len(title))}")

    sep()
    print("  ARC 키워드 분석 결과")
    print(f"  키워드: {kws}   |   모드: {mode}   |   {date}")
    sep()

    # ── C. Coverage ──────────────────────────────────────────────
    coverage = result.get("C_coverage", [])
    if coverage:
        section("키워드 적합도")
        for c in coverage:
            pct     = c.get("coverage_percent", 0)
            bar_len = min(int(pct / 5), 20)
            bar     = "█" * bar_len + "░" * (20 - bar_len)
            h = c.get("high_count", 0)
            m = c.get("medium_count", 0)
            lo = c.get("low_count", 0)
            print(f"  {c['keyword']:<14} {bar} {pct:5.1f}%   H:{h} M:{m} L:{lo}")

    # ── D. 매칭 경험 ─────────────────────────────────────────────
    matched = result.get("D_matched_experiences", [])
    if matched:
        section("매칭된 경험")
        REL_ICON  = {"high": "●", "medium": "◑", "low": "○"}
        CONF_ICON = {"high": "★", "medium": "☆", "low": "△"}
        for kw_block in matched:
            kw   = kw_block.get("keyword", "")
            exps = kw_block.get("experiences", [])
            if not exps:
                continue
            print(f"\n  ▶ {kw}  ({len(exps)}개)")
            print("     관련도: ●high  ◑medium  ○low  |  확신도: ★high  ☆medium  △low")
            for exp in exps:
                rel     = exp.get("relevance", "medium")
                conf    = exp.get("confidence", "medium")
                r_icon  = REL_ICON.get(rel, "·")
                c_icon  = CONF_ICON.get(conf, "·")
                title   = exp.get("career_title", "")
                is_ref  = exp.get("is_reference_only", False)
                ref_tag = "  [참고용]" if is_ref else ""
                rel_sum = exp.get("relevance_summary", "")
                print(f"    {r_icon} {title}{ref_tag}  (확신도:{c_icon})")
                if rel_sum:
                    print(f"      → {rel_sum}")
                evs = exp.get("evidence", [])
                if evs:
                    q = evs[0].get("source_quote", "")
                    if q:
                        print(f'         "{q[:72]}{"..." if len(q) > 72 else ""}')

    # ── E. 스토리라인 ─────────────────────────────────────────────
    storylines = result.get("E_storylines", [])
    if storylines:
        section("스토리라인")
        for sl in storylines:
            kw      = sl.get("keyword", "")
            title   = sl.get("storyline_title", "")
            tagline = sl.get("tagline", "")
            narr    = sl.get("narrative", "")
            tps     = sl.get("turning_points", [])
            cl      = sl.get("connective_logic", [])
            struct  = sl.get("structure", {})

            print(f"\n  ▶ [{kw}]  {title}")
            if tagline:
                print(f"     💬 {tagline}")

            tl_status = sl.get("timeline_status", "")
            tl_note   = sl.get("timeline_note", "")
            chrono    = sl.get("chronological_sequence", [])
            if tl_status:
                tl_icon = {"시간순_확인됨": "🕐", "일부_불명확": "⚠️", "대부분_불명확": "❓"}.get(tl_status, "")
                print(f"     {tl_icon} 시간 순서: {tl_status}")
                if tl_note:
                    print(f"       주의: {tl_note[:100]}")
            if chrono:
                print("     [ 경험 순서 ]  " + "  →  ".join(
                    f"{c['experience']}({'시기미상' if not c.get('is_dated') else c.get('period','')})"
                    for c in chrono
                ))

            if narr:
                print("\n     [ 서사 ]")
                words, line, out_lines = narr.split(), "", []
                for w in words:
                    if len(line) + len(w) + 1 > 66:
                        out_lines.append(line)
                        line = w
                    else:
                        line = (line + " " + w).strip()
                if line:
                    out_lines.append(line)
                for ln in out_lines:
                    print(f"     {ln}")

            if tps:
                print("\n     [ 전환점 ]")
                for tp in tps:
                    exp_name = tp.get("experience", "")
                    trigger  = tp.get("trigger", "")[:80]
                    changed  = tp.get("what_changed", "")[:80]
                    print(f"     ⟳ {exp_name}")
                    if trigger:
                        print(f"       촉발: {trigger}")
                    if changed:
                        print(f"       변화: {changed}")

            if cl:
                print("\n     [ 연결 고리 ]")
                for c in cl:
                    rel_type = c.get("relation_type", "")
                    conn     = c.get("connection", "")[:80]
                    fr       = c.get("from_experience", "")
                    to       = c.get("to_experience", "")
                    print(f"     {fr}  →[{rel_type}]→  {to}")
                    if conn:
                        print(f"       {conn}")

            if struct.get("destination"):
                print(f"\n     [ 현재 역량 ]")
                print(f"     {struct['destination'][:100]}")

    # ── F. 개선 방향 ─────────────────────────────────────────────
    guide        = result.get("F_improvement_guide", {})
    direction    = guide.get("overall_direction", {})
    enhancements = guide.get("information_enhancement", [])
    expansions   = guide.get("experience_expansion", [])
    _prio_order  = {"높음": 0, "중간": 1, "낮음": 2}
    _prio_icon   = {"높음": "🔴", "중간": "🟡", "낮음": "⚪"}

    if direction:
        section("종합 방향성")
        if direction.get("current_profile_summary"):
            print(f"  [진단] {direction['current_profile_summary']}")
        if direction.get("short_term"):
            print(f"\n  ▸ 단기 (3개월): {direction['short_term']}")
        if direction.get("mid_term"):
            print(f"  ▸ 중기 (1년) : {direction['mid_term']}")
        if direction.get("priority_keyword"):
            reason = direction.get("priority_reason", "")
            print(f"\n  ★ 최우선 보강 키워드: {direction['priority_keyword']}")
            if reason:
                print(f"    이유: {reason}")

    if enhancements or expansions:
        section("보완 제안 (우선순위순)")
        for e in sorted(enhancements, key=lambda x: _prio_order.get(x.get("priority", "중간"), 1))[:3]:
            icon = _prio_icon.get(e.get("priority", "중간"), "·")
            print(f"  {icon} [기록 보강] {e.get('target', '')} — {e.get('missing', '')}")
            if e.get("how_to_add"):
                print(f"    → {e['how_to_add']}")
        for e in sorted(expansions, key=lambda x: _prio_order.get(x.get("priority", "중간"), 1))[:2]:
            icon = _prio_icon.get(e.get("priority", "중간"), "·")
            print(f"  {icon} [경험 확장] {e.get('gap_description', '')} → {e.get('suggested_experience_type', '')}")

    sep()
    print("  전체 JSON: result 변수  |  스토리라인만: result['E_storylines']")
    sep()


# ══════════════════════════════════════════════════════════════════
# ▼▼▼  여기만 수정하세요  ▼▼▼
# ══════════════════════════════════════════════════════════════════

CAREER_INPUT = """
여기에 경력 데이터를 입력하세요.
URL, 텍스트, 복합 입력 모두 가능합니다.
"""

KEYWORDS = ["키워드1", "키워드2"]  # 개수 제한 없음
TARGET   = ""                       # 목표 직무/시나리오 (없으면 빈 문자열)


# ══════════════════════════════════════════════════════════════════
# 10. Main (엔트리포인트)
# ══════════════════════════════════════════════════════════════════

def main(
    career_input: str,
    keywords: list[str],
    target: str,
) -> dict:
    """
    전체 파이프라인 실행 엔트리포인트.
    경력 파싱 → 키워드 분석 → 요약 출력을 순서대로 수행합니다.

    Args:
        career_input: 경력 데이터 (URL/파일경로/텍스트/혼합).
                      None이면 상단 CAREER_INPUT 사용.
        keywords:     분석 키워드 목록. None이면 상단 KEYWORDS 사용.
        target:       목표 직무/시나리오. None이면 상단 TARGET 사용.

    Returns:
        {"careers": [...], "result": {...}} — 파싱된 경력과 분석 결과.
        입력이 placeholder 상태이거나 파싱 실패 시 {"careers": [], "result": {}}.

    사용 예:
        # 1) 상단 설정값으로 실행 (Colab 셀 실행과 동일)
        output = main()

        # 2) 인자 직접 전달
        output = main(
            career_input="https://my-portfolio.notion.site/...",
            keywords=["리더십", "분석력"],
            target="자산운용사 퀀트 애널리스트",
        )
        careers, result = output["careers"], output["result"]
    """
    career_input = CAREER_INPUT if career_input is None else career_input
    keywords     = KEYWORDS     if keywords     is None else keywords
    target       = TARGET       if target       is None else target

    # placeholder 미수정 상태 감지 → 안내 후 종료
    if "여기에 경력 데이터를 입력하세요" in career_input or "키워드1" in keywords:
        print("=" * 60)
        print("  ⚠️  CAREER_INPUT / KEYWORDS를 먼저 수정하세요.")
        print("  추천 키워드 목록은 show_keywords()로 확인할 수 있습니다.")
        print("  또는 main(career_input=..., keywords=[...], target=...)로")
        print("  인자를 직접 전달해 실행할 수 있습니다.")
        print("=" * 60)
        return {"careers": [], "result": {}}

    # 1) 경력 파싱 (+ 임베딩 사전 생성)
    careers = parse_careers(career_input)
    if not careers:
        print("[main] 경력 파싱 결과가 비어 있어 분석을 중단합니다.")
        return {"careers": [], "result": {}}

    # 2) 키워드 분석 (경험 수에 따라 LLM/K-NN 자동 라우팅)
    result = analyze_keywords(keywords, careers, target=target)

    # 3) 요약 출력
    print_result_summary(result)

    return {"careers": careers, "result": result, "status": "success"}


if __name__ == "__main__":
    output  = main()
    careers = output["careers"]   # 이후 셀에서 재사용 가능
    result  = output["result"]